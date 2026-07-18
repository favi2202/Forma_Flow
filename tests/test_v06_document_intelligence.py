from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient
from openpyxl import Workbook
from reportlab.pdfgen import canvas

from app import (
    app,
    classify_document,
    dataset_from_rows,
    parse_uploaded_file,
)

client = TestClient(app)


def docx_bytes(builder) -> bytes:
    document = Document()
    builder(document)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def xlsx_bytes(builder) -> bytes:
    workbook = Workbook()
    builder(workbook)
    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def pdf_bytes(lines: list[str]) -> bytes:
    output = BytesIO()
    pdf = canvas.Canvas(output)
    y = 800
    for line in lines:
        pdf.drawString(50, y, line)
        y -= 22
    pdf.save()
    return output.getvalue()


def test_numeric_monitoring_data_row_is_not_selected_as_header():
    rows = [
        ["", "Sinflar", "Texnologiya", "Musiqa", "Tasviriy san'at", "Jismoniy tarbiya", "Sinflar kesimida"],
        [1, "2-A", 72, 65, 60, 67, 66],
        [2, "2-B", 70, 68, 64, 64, 67],
        [3, "3-A", 74, 70, 70, 70, 71],
    ]
    dataset = dataset_from_rows(rows, "monitoring.docx", "Table 1", "monitoring")
    assert dataset is not None
    assert dataset["header_index"] == 0
    assert dataset["dataset_type"] == "monitoring_table"
    assert {column.key for column in dataset["columns"]} >= {"row_number", "class"}
    assert len(dataset["records"]) == 3


def test_single_column_roster_generates_number_name_and_class():
    def build(workbook: Workbook) -> None:
        sheet = workbook.active
        sheet.title = "2-B"
        sheet.append(["2-B sinf o'quvchilar ro'yxati"])
        sheet.append(["O'quvchining ism va familiyasi"])
        sheet.append(["Anvar Aliyev"])
        sheet.append(["Malika Saidova"])
        sheet.append(["Sardor Karimov"])

    parsed = parse_uploaded_file("2-B roster.xlsx", xlsx_bytes(build))
    assert parsed["document_type"] == "student_roster_document"
    assert len(parsed["datasets"]) == 1
    dataset = parsed["datasets"][0]
    assert dataset["dataset_type"] == "student_roster"
    assert len(dataset["records"]) == 3
    assert {column.key for column in dataset["columns"]} >= {"row_number", "student_name", "class"}
    assert {record["class"] for record in dataset["records"]} == {"2-B"}


def test_methodical_document_does_not_create_fake_student_roster():
    def build(document: Document) -> None:
        document.add_heading("METODIK TAVSIYA", 0)
        document.add_paragraph("Annotatsiya")
        document.add_paragraph("Ushbu metodik qo'llanma interfaol metodlar va dars jarayoni haqida.")
        table = document.add_table(rows=5, cols=1)
        for index, text in enumerate(["MUNDARIJA", "Kirish", "Asosiy qism", "Xulosa", "Adabiyotlar"]):
            table.cell(index, 0).text = text

    parsed = parse_uploaded_file("metodik_tavsiya.docx", docx_bytes(build))
    assert parsed["document_type"] == "methodical_document"
    assert parsed["datasets"] == []


def test_financial_document_is_classified_and_not_called_roster():
    def build(document: Document) -> None:
        document.add_heading("АКТ СВЕРКА", 0)
        document.add_paragraph("Сальдо, дебет и кредит")
        table = document.add_table(rows=3, cols=4)
        for col, value in enumerate(["Дата", "Документ", "Дебет", "Кредит"]):
            table.cell(0, col).text = value
        for col, value in enumerate(["01.01.2026", "A-1", "100", "0"]):
            table.cell(1, col).text = value
        for col, value in enumerate(["02.01.2026", "A-2", "0", "50"]):
            table.cell(2, col).text = value

    parsed = parse_uploaded_file("akt_sverka.docx", docx_bytes(build))
    assert parsed["document_type"] == "financial_document"
    assert all(dataset["dataset_type"] != "student_roster" for dataset in parsed["datasets"])


def test_dominant_promotion_tables_classify_document_as_promotion():
    promotion_columns = []
    for key, label in [
        ("row_number", "No."),
        ("class", "Class"),
        ("student_name", "Student name"),
        ("promoted_class", "Promoted class"),
    ]:
        from app import ColumnInfo
        promotion_columns.append(ColumnInfo(key, label, label, 0.99, "alias"))

    roster_columns = promotion_columns[:3]
    datasets = []
    for index in range(8):
        datasets.append({
            "dataset_type": "promotion_table",
            "columns": promotion_columns,
            "records": [{"row_number": i, "class": "4-A", "student_name": f"Student {i}", "promoted_class": "5-A"} for i in range(1, 11)],
        })
    datasets.append({
        "dataset_type": "student_roster",
        "columns": roster_columns,
        "records": [{"row_number": 1, "class": "4-A", "student_name": "Student One"}],
    })

    result = classify_document("jadval.docx", "O'quvchi jadvallari", datasets)
    assert result["type"] == "promotion_document"


def test_mixed_docx_creates_separate_dataset_groups_and_can_switch():
    def build(document: Document) -> None:
        document.add_heading("School data", 0)
        roster = document.add_table(rows=3, cols=3)
        for col, value in enumerate(["№", "O'quvchi F.I.Sh.", "Sinf"]):
            roster.cell(0, col).text = value
        for row, values in enumerate([[1, "Ali Karimov", "5-A"], [2, "Malika Saidova", "5-A"]], start=1):
            for col, value in enumerate(values):
                roster.cell(row, col).text = str(value)

        monitoring = document.add_table(rows=3, cols=5)
        for col, value in enumerate(["№", "Ism sharifi", "1", "2", "Jami"]):
            monitoring.cell(0, col).text = value
        for row, values in enumerate([[1, "Ali Karimov", 1, 1, 2], [2, "Malika Saidova", 1, 0, 1]], start=1):
            for col, value in enumerate(values):
                monitoring.cell(row, col).text = str(value)

    response = client.post(
        "/api/upload",
        files=[("files", ("mixed.docx", docx_bytes(build), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))],
    )
    assert response.status_code == 200, response.text
    payload = response.json()
    assert len(payload["dataset_groups"]) == 2
    types = {group["dataset_type"] for group in payload["dataset_groups"]}
    assert types == {"student_roster", "monitoring_table"}

    monitoring_group = next(group for group in payload["dataset_groups"] if group["dataset_type"] == "monitoring_table")
    selected = client.post(
        "/api/select-dataset",
        json={"session_id": payload["session_id"], "dataset_id": monitoring_group["id"]},
    )
    assert selected.status_code == 200
    selected_payload = selected.json()
    assert selected_payload["dataset_type"] == "monitoring_table"
    assert any(column["key"].startswith("question_") for column in selected_payload["columns"])


def test_empty_monitoring_template_is_not_exposed_as_dataset():
    def build(workbook: Workbook) -> None:
        sheet = workbook.active
        sheet.title = "Empty template"
        sheet.append(["№", "Ism sharifi", 1, 2, 3, 4, 5, "Jami"])
        for number in range(1, 20):
            sheet.append([number, "", "", "", "", "", "", ""])

    parsed = parse_uploaded_file("monitoring_template.xlsx", xlsx_bytes(build))
    assert parsed["datasets"] == []
    assert parsed["warnings"]


def test_pdf_assessment_material_is_classified_without_fake_table():
    content = pdf_bytes([
        "7-DARS TEST SAVOLLARI",
        "1. Qaysi javob to'g'ri?",
        "A) Birinchi B) Ikkinchi C) Uchinchi D) To'rtinchi",
        "2. Qaysi gapda fe'l bor?",
        "3. Qaysi javob noto'g'ri?",
        "4. Savol matni",
        "5. Savol matni",
    ])
    parsed = parse_uploaded_file("7-dars.pdf", content)
    assert parsed["document_type"] == "assessment_material"
    assert parsed["datasets"] == []


def test_meeting_minutes_remain_meeting_even_with_large_promotion_appendix():
    from app import ColumnInfo

    columns = [
        ColumnInfo("row_number", "No.", "No.", 0.99, "alias"),
        ColumnInfo("class", "Class", "Class", 0.99, "alias"),
        ColumnInfo("student_name", "Student name", "Student name", 0.99, "alias"),
        ColumnInfo("promoted_class", "Promoted class", "Promoted class", 0.99, "alias"),
    ]
    datasets = [{
        "dataset_type": "promotion_table",
        "columns": columns,
        "records": [{"row_number": i, "class": "5-A", "student_name": f"Student {i}", "promoted_class": "6-A"} for i in range(1, 401)],
    }]
    text = "Ped kengash yig'ilishi bayonnomasi. KUN TARTIBI. Yig'ilish qaror qiladi. Monitoring natijalari."
    result = classify_document("PED KENGASH.docx", text, datasets)
    assert result["type"] == "meeting_minutes"
