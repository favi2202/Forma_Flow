from io import BytesIO
from pathlib import Path

import xlwt
from fastapi.testclient import TestClient
from openpyxl import Workbook, load_workbook

from app import app, next_class_value, rows_to_records, is_repeated_header_row

client = TestClient(app)


def make_xlsx() -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "8A"
    ws.append(["Class report 2026"])
    ws.append([])
    ws.append(["F.I.O.", "Sinf", "Baho", "Telefon"])
    ws.append(["Ali Karimov", "8A", 5, "+998901111111"])
    ws.append(["Malika Saidova", "8A", 4, "+998902222222"])
    out = BytesIO()
    wb.save(out)
    return out.getvalue()


def make_xls() -> bytes:
    wb = xlwt.Workbook()
    ws = wb.add_sheet("9A")
    values = [
        ["O‘quvchi ismi", "Class", "Mark", "Address"],
        ["Sardor Ergashev", "9A", 5, "Tashkent"],
        ["Madina Sobirova", "9A", 4, "Tashkent"],
    ]
    for r, row in enumerate(values):
        for c, value in enumerate(row):
            ws.write(r, c, value)
    out = BytesIO()
    wb.save(out)
    return out.getvalue()


def test_health():
    response = client.get("/health")
    assert response.status_code == 200
    assert response.json()["status"] == "ok"


def test_upload_and_exports():
    response = client.post(
        "/api/upload",
        files=[
            ("files", ("class8.xlsx", make_xlsx(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")),
            ("files", ("class9.xls", make_xls(), "application/vnd.ms-excel")),
        ],
    )
    assert response.status_code == 200, response.text
    payload = response.json()
    assert payload["row_count"] == 4
    keys = {column["key"] for column in payload["columns"]}
    assert {"student_name", "class", "grade"}.issubset(keys)

    export_request = {
        "session_id": payload["session_id"],
        "columns": [
            {"key": "student_name", "name": "Student name"},
            {"key": "class", "name": "Class"},
            {"key": "grade", "name": "Grade"},
        ],
        "fixed_columns": [{"name": "Academic year", "value": "2026–2027"}],
        "format": "xlsx",
    }
    xlsx_response = client.post("/api/export", json=export_request)
    assert xlsx_response.status_code == 200
    wb = load_workbook(BytesIO(xlsx_response.content), data_only=True)
    ws = wb.active
    assert ws.max_row == 5
    assert ws["A2"].value == "Ali Karimov"
    assert ws["D5"].value == "2026–2027"

    for fmt in ("csv", "docx"):
        export_request["format"] = fmt
        result = client.post("/api/export", json=export_request)
        assert result.status_code == 200
        assert len(result.content) > 100


def make_html_xls() -> bytes:
    return b"""<html><body><table>
    <tr><th>F.I.O.</th><th>Sinf</th><th>Baho</th></tr>
    <tr><td>Ali Karimov</td><td>8A</td><td>5</td></tr>
    </table></body></html>"""


def make_spreadsheetml_xls() -> bytes:
    return b"""<?xml version=\"1.0\"?>
    <Workbook xmlns=\"urn:schemas-microsoft-com:office:spreadsheet\"
      xmlns:ss=\"urn:schemas-microsoft-com:office:spreadsheet\">
      <Worksheet ss:Name=\"Class\"><Table>
        <Row><Cell><Data ss:Type=\"String\">F.I.O.</Data></Cell><Cell><Data ss:Type=\"String\">Sinf</Data></Cell></Row>
        <Row><Cell><Data ss:Type=\"String\">Malika Saidova</Data></Cell><Cell><Data ss:Type=\"String\">8A</Data></Cell></Row>
      </Table></Worksheet>
    </Workbook>"""


def test_html_saved_as_xls_is_parsed():
    response = client.post(
        "/api/upload",
        files=[("files", ("school_export.xls", make_html_xls(), "application/vnd.ms-excel"))],
    )
    assert response.status_code == 200, response.text
    payload = response.json()
    assert payload["row_count"] == 1
    assert "HTML table" in payload["files"][0]["detected_type"]


def test_spreadsheetml_saved_as_xls_is_parsed():
    response = client.post(
        "/api/upload",
        files=[("files", ("school_export.xls", make_spreadsheetml_xls(), "application/vnd.ms-excel"))],
    )
    assert response.status_code == 200, response.text
    payload = response.json()
    assert payload["row_count"] == 1
    assert "SpreadsheetML" in payload["files"][0]["detected_type"]


def test_invalid_xls_returns_specific_error():
    response = client.post(
        "/api/upload",
        files=[("files", ("broken.xls", b"not really an excel file", "application/vnd.ms-excel"))],
    )
    assert response.status_code == 400
    detail = response.json()["detail"]
    assert detail["files"][0]["name"] == "broken.xls"
    assert "detected content" in detail["files"][0]["error"].lower()



def test_next_class_value_rules():
    assert next_class_value("5-b") == "6-b"
    assert next_class_value("8-A") == "9-A"
    assert next_class_value("9-a") == ""
    assert next_class_value("10-Б") == "11-Б"
    assert next_class_value("11") == ""
    assert next_class_value("not a class") == ""


def test_preview_and_export_with_derived_next_class():
    response = client.post(
        "/api/upload",
        files=[("files", ("class8.xlsx", make_xlsx(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"))],
    )
    assert response.status_code == 200, response.text
    payload = response.json()

    request = {
        "session_id": payload["session_id"],
        "columns": [
            {"key": "student_name", "name": "Student name"},
            {"key": "class", "name": "Class"},
        ],
        "fixed_columns": [],
        "derived_columns": [
            {
                "name": "Next class",
                "kind": "next_class",
                "source_key": "class",
                "stop_grades": [9, 11],
                "amount": 1,
                "start": 1,
            },
            {
                "name": "No.",
                "kind": "sequence",
                "source_key": None,
                "stop_grades": [9, 11],
                "amount": 1,
                "start": 1,
            },
        ],
        "options": {
            "trim_whitespace": True,
            "remove_duplicates": False,
            "skip_blank_key": None,
            "sort_key": "student_name",
        },
    }

    preview_response = client.post("/api/preview", json={**request, "limit": 50})
    assert preview_response.status_code == 200, preview_response.text
    preview = preview_response.json()
    assert preview["row_count"] == 2
    assert preview["headers"] == ["Student name", "Class", "Next class", "No."]
    assert preview["rows"][0]["Next class"] == "9A"
    assert preview["rows"][0]["No."] == 1

    export_response = client.post("/api/export", json={**request, "format": "xlsx"})
    assert export_response.status_code == 200, export_response.text
    workbook = load_workbook(BytesIO(export_response.content), data_only=True)
    sheet = workbook.active
    assert sheet["C2"].value == "9A"
    assert sheet["D2"].value == 1


def test_repeated_secondary_header_rows_are_removed():
    rows = [
        ["№", "O‘quvchi F.I.Sh.", "Sinf"],
        ["Telefon", "Elektron pochta", ""],
        [1, "Abdumalikova Marjona", "5-E"],
        [2, "Barotova Lola", "5-E"],
        ["№", "O‘quvchi F.I.Sh.", "Sinf"],
        ["Telefon", "Elektron pochta", ""],
        [1, "Abdumalikova Sevara", "6-A"],
    ]

    _, columns, records = rows_to_records(rows, "school.xls", "Sheet1")

    assert len(records) == 3
    student_key = next(column.key for column in columns if column.key == "student_name")
    class_key = next(column.key for column in columns if column.key == "class")
    assert [record[student_key] for record in records] == [
        "Abdumalikova Marjona",
        "Barotova Lola",
        "Abdumalikova Sevara",
    ]
    assert [record[class_key] for record in records] == ["5-E", "5-E", "6-A"]


def test_header_cleanup_is_conservative_for_student_rows():
    primary = ["№", "O‘quvchi F.I.Sh.", "Sinf"]
    assert is_repeated_header_row(["Telefon", "Elektron pochta", ""], primary)
    assert is_repeated_header_row(["№", "O‘quvchi F.I.Sh.", "Sinf"], primary)
    assert not is_repeated_header_row([1, "Ali Karimov", "5-E"], primary)
    assert not is_repeated_header_row([2, "Malika Saidova", "6-A"], primary)


def make_sectioned_xls_with_repeated_headers() -> bytes:
    wb = xlwt.Workbook()
    ws = wb.add_sheet("Classes")
    values = [
        ["№", "O‘quvchi F.I.Sh.", "Sinf"],
        ["Telefon", "Elektron pochta", ""],
        [1, "Abdumalikova Marjona", "5-E"],
        [2, "Barotova Lola", "5-E"],
        ["№", "O‘quvchi F.I.Sh.", "Sinf"],
        ["Telefon", "Elektron pochta", ""],
        [1, "Abdumalikova Sevara", "6-A"],
    ]
    for row_index, row in enumerate(values):
        for column_index, value in enumerate(row):
            ws.write(row_index, column_index, value)
    output = BytesIO()
    wb.save(output)
    return output.getvalue()


def test_upload_pipeline_removes_repeated_headers_before_export():
    response = client.post(
        "/api/upload",
        files=[
            (
                "files",
                (
                    "sectioned.xls",
                    make_sectioned_xls_with_repeated_headers(),
                    "application/vnd.ms-excel",
                ),
            )
        ],
    )
    assert response.status_code == 200, response.text
    payload = response.json()
    assert payload["row_count"] == 3

    request = {
        "session_id": payload["session_id"],
        "columns": [
            {"key": "student_name", "name": "O‘quvchi F.I.Sh."},
            {"key": "class", "name": "Sinf"},
        ],
        "fixed_columns": [],
        "derived_columns": [
            {
                "name": "Tartib raqam",
                "kind": "sequence",
                "source_key": None,
                "stop_grades": [9, 11],
                "amount": 1,
                "start": 1,
            }
        ],
        "options": {
            "trim_whitespace": True,
            "remove_duplicates": False,
            "skip_blank_key": None,
            "sort_key": None,
        },
        "format": "xlsx",
    }
    export_response = client.post("/api/export", json=request)
    assert export_response.status_code == 200, export_response.text
    workbook = load_workbook(BytesIO(export_response.content), data_only=True)
    sheet = workbook.active
    exported_names = [sheet.cell(row=row, column=1).value for row in range(2, sheet.max_row + 1)]
    assert exported_names == [
        "Abdumalikova Marjona",
        "Barotova Lola",
        "Abdumalikova Sevara",
    ]
    assert [sheet.cell(row=row, column=3).value for row in range(2, sheet.max_row + 1)] == [1, 2, 3]


def make_docx_tables(tables: list[list[list[object]]]) -> bytes:
    from docx import Document

    document = Document()
    for table_index, rows in enumerate(tables):
        if table_index:
            document.add_paragraph()
        column_count = max(len(row) for row in rows)
        table = document.add_table(rows=0, cols=column_count)
        for source_row in rows:
            cells = table.add_row().cells
            for index, value in enumerate(source_row):
                cells[index].text = str(value)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def test_docx_single_table_is_parsed():
    content = make_docx_tables([
        [
            ["F.I.O.", "Sinf", "Baho"],
            ["Ali Karimov", "5-A", 5],
            ["Malika Saidova", "5-A", 4],
        ]
    ])
    response = client.post(
        "/api/upload",
        files=[
            (
                "files",
                (
                    "class.docx",
                    content,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            )
        ],
    )
    assert response.status_code == 200, response.text
    payload = response.json()
    assert payload["row_count"] == 2
    summary = payload["files"][0]
    assert summary["source_kind"] == "word"
    assert summary["table_count"] == 1
    assert summary["tables_differ"] is False


def test_similar_docx_tables_are_merged():
    content = make_docx_tables([
        [
            ["F.I.O.", "Sinf"],
            ["Ali Karimov", "5-A"],
        ],
        [
            ["F.I.O.", "Sinf"],
            ["Malika Saidova", "6-B"],
            ["Sardor Ergashev", "6-B"],
        ],
    ])
    response = client.post(
        "/api/upload",
        files=[("files", ("classes.docx", content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))],
    )
    assert response.status_code == 200, response.text
    payload = response.json()
    assert payload["row_count"] == 3
    summary = payload["files"][0]
    assert summary["table_count"] == 2
    assert summary["tables_merged"] is True
    assert summary["warning_code"] == "similar_word_tables_merged"


def test_different_docx_tables_are_flagged_and_best_table_is_used():
    content = make_docx_tables([
        [
            ["Phone", "Email"],
            ["+998900000001", "one@example.com"],
        ],
        [
            ["F.I.O.", "Sinf", "Baho"],
            ["Ali Karimov", "5-A", 5],
            ["Malika Saidova", "5-A", 4],
            ["Sardor Ergashev", "5-A", 5],
        ],
    ])
    response = client.post(
        "/api/upload",
        files=[("files", ("mixed.docx", content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))],
    )
    assert response.status_code == 200, response.text
    payload = response.json()
    assert payload["row_count"] == 3
    summary = payload["files"][0]
    assert summary["tables_differ"] is True
    assert summary["selected_table"] == 2
    assert summary["warning_code"] == "different_word_tables"
    keys = {column["key"] for column in payload["columns"]}
    assert {"student_name", "class", "grade"}.issubset(keys)
