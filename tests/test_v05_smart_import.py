from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient

from app import app, infer_class, parse_docx, parse_html_table, recognize_header

client = TestClient(app)


def make_headerless_roster_docx(title: str = '11-"B" sinf o\'quvchilari ro\'yxati') -> bytes:
    document = Document()
    document.add_paragraph(title)
    table = document.add_table(rows=0, cols=3)
    for number, name in [(1, "Ali Karimov"), (2, "Malika Saidova")]:
        cells = table.add_row().cells
        cells[0].text = str(number)
        cells[1].text = name
        cells[2].text = ""
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def make_two_level_html() -> bytes:
    return '''<html><body><table>
    <tr><th rowspan="2">№</th><th rowspan="2">O‘quvchi</th><th rowspan="2">Sinf</th><th rowspan="2">Qarindoshlar</th><th colspan="2">Qarindoshlarning aloqa ma'lumotlari</th></tr>
    <tr><th>Telefon</th><th>Elektron pochta</th></tr>
    <tr><td>1</td><td>Ali Karimov</td><td>9-A</td><td rowspan="2">Karimov A. (Ota)</td><td>998901111111</td><td>parent1@example.com</td></tr>
    <tr><td>1</td><td>Ali Karimov</td><td>9-A</td><td>998902222222</td><td>parent2@example.com</td></tr>
    <tr><td>2</td><td>Malika Saidova</td><td>9-A</td><td>Saidova M. (Ona)</td><td>998903333333</td><td>parent3@example.com</td></tr>
    </table></body></html>'''.encode('utf-8')


def test_short_order_heading_is_not_student_name():
    assert recognize_header("T/N").key == "row_number"
    assert recognize_header("№").key == "row_number"
    assert recognize_header("O`quvchining ismi,sharifi").key == "student_name"


def test_headerless_word_roster_keeps_first_student_and_infers_class():
    result = parse_docx(make_headerless_roster_docx(), "11-b.docx")
    assert result["header_mode"] == "inferred_roster"
    assert len(result["records"]) == 2
    assert result["records"][0]["student_name"] == "Ali Karimov"
    assert result["records"][0]["row_number"] == "1"
    assert result["records"][0]["class"] == "11-B"
    assert {column.key for column in result["columns"]} == {"row_number", "student_name", "class"}


def test_two_level_html_header_and_continuation_contacts_are_merged():
    result = parse_html_table(make_two_level_html(), "9-a.xls")
    keys = {column.key for column in result["columns"]}
    assert {"row_number", "student_name", "class", "parent_name", "parent_phone", "parent_email"}.issubset(keys)
    assert result["header_mode"] == "multi_row"
    assert result["continuation_rows_merged"] == 1
    assert len(result["records"]) == 2
    first = result["records"][0]
    assert first["student_name"] == "Ali Karimov"
    assert first["parent_phone"] == "998901111111 | 998902222222"
    assert first["parent_email"] == "parent1@example.com | parent2@example.com"


def test_class_inference_prefers_document_title_over_conflicting_sheet_name():
    info = infer_class([
        ('10-,,B" sinf o`quvchilari', "document title", 100),
        ("10-B SINF.xls", "file name", 80),
        ("11-B sinf", "sheet name", 40),
    ])
    assert info["value"] == "10-B"
    assert info["source"] == "document title"
    assert info["conflict"] is True


def test_upload_summary_exposes_smart_import_diagnostics():
    response = client.post(
        "/api/upload",
        files=[
            (
                "files",
                (
                    "11-b.docx",
                    make_headerless_roster_docx(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            ),
            ("files", ("9-a.xls", make_two_level_html(), "application/vnd.ms-excel")),
        ],
    )
    assert response.status_code == 200, response.text
    payload = response.json()
    assert payload["row_count"] == 4
    assert payload["files"][0]["header_mode"] == "inferred_roster"
    assert payload["files"][0]["inferred_class"] == "11-B"
    assert payload["files"][1]["header_mode"] == "multi_row"
    assert payload["files"][1]["continuation_rows_merged"] == 1
    selected = {column["key"] for column in payload["columns"] if column["default_selected"]}
    assert selected == {"row_number", "student_name", "class"}
    sensitive = {column["key"] for column in payload["columns"] if column["sensitive"]}
    assert {"parent_phone", "parent_email"}.issubset(sensitive)
