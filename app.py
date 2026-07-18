from __future__ import annotations

import csv
import io
import json
import re
import secrets
import shutil
import subprocess
import tempfile
import unicodedata
import zipfile
import xml.etree.ElementTree as ET
from collections import OrderedDict
from dataclasses import dataclass
from datetime import date, datetime
from difflib import SequenceMatcher
from pathlib import Path
from html import unescape
from typing import Any, Literal

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.responses import FileResponse, HTMLResponse, Response, StreamingResponse
from fastapi.staticfiles import StaticFiles
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from bs4 import BeautifulSoup
from pydantic import BaseModel, Field

try:
    import xlrd
except ImportError:  # handled with a clear message at runtime
    xlrd = None

ROOT = Path(__file__).resolve().parent
STATIC_DIR = ROOT / "static"
MAX_FILE_BYTES = 25 * 1024 * 1024
MAX_ROWS_PER_FILE = 50_000
MAX_FILES = 50

app = FastAPI(title="FormaFlow Local", version="0.6.0")
app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")


@dataclass
class ColumnInfo:
    key: str
    label: str
    original: str
    confidence: float
    method: str


SESSIONS: dict[str, dict[str, Any]] = {}

FIELD_DEFINITIONS: dict[str, dict[str, Any]] = {
    "row_number": {
        "label": "No.",
        "aliases": [
            "no", "no.", "number", "№", "n", "n.", "t/n", "t n", "tn",
            "tartib raqam", "tartib nomer", "tartib soni", "ketma ket raqam",
            "п/п", "номер", "№ п/п", "порядковый номер", "sequence", "order",
        ],
    },
    "student_name": {
        "label": "Student name",
        "aliases": [
            "student name", "student", "pupil", "full name", "name surname",
            "fio", "f i o", "f.i.o", "фио", "ф и о", "ф.и.о",
            "ученик", "имя ученика", "фамилия имя", "фамилия и имя",
            "oquvchi", "o'quvchi", "o‘quvchi", "oquvchi ismi",
            "o'quvchi ismi", "o‘quvchi ismi", "oquvchi f i sh",
            "o'quvchi f.i.sh.", "o‘quvchi f.i.sh.", "fish", "f i sh", "f.i.sh",
            "familiya ism sharif", "familiya ismi", "ism familiyasi",
            "o'quvchilarning ism familiyasi", "o‘quvchilarning ism familiyasi",
            "o'quvchining ismi sharifi", "o‘quvchining ismi sharifi",
            "o`quvchining ismi,sharifi", "o'quvchining ismi, sharifi",
        ],
    },
    "first_name": {"label": "First name", "aliases": ["first name", "ism", "имя"]},
    "last_name": {"label": "Last name", "aliases": ["last name", "surname", "familiya", "фамилия"]},
    "class": {"label": "Class", "aliases": ["class", "grade class", "sinf", "sinfi", "класс"]},
    "grade": {"label": "Grade", "aliases": ["grade", "mark", "score", "result", "baho", "балл", "оценка"]},
    "age": {"label": "Age", "aliases": ["age", "yosh", "возраст"]},
    "gender": {"label": "Gender", "aliases": ["gender", "sex", "jinsi", "пол"]},
    "birth_date": {
        "label": "Birth date",
        "aliases": ["birth date", "date of birth", "birthday", "dob", "tugilgan sana", "tug'ilgan sana", "tug‘ilgan sana", "дата рождения"],
    },
    "pinfl": {"label": "JSHSHIR", "aliases": ["jshshir", "pinfl", "пинфл"]},
    "email": {"label": "Email", "aliases": ["email", "e mail", "mail", "elektron pochta", "e pochta", "электронная почта", "почта"]},
    "phone": {"label": "Phone", "aliases": ["phone", "phone number", "telephone", "tel", "telefon", "телефон", "номер телефона"]},
    "address": {"label": "Address", "aliases": ["address", "home address", "manzil", "адрес", "домашний адрес"]},
    "parent_name": {
        "label": "Relative / guardian",
        "aliases": ["parent name", "guardian", "relative", "qarindoshlar", "ota ona", "родитель", "фио родителя", "родственники"],
    },
    "parent_phone": {
        "label": "Relative phone",
        "aliases": [
            "parent phone", "guardian phone", "relative phone", "ota ona telefoni", "телефон родителя",
            "qarindoshlarning aloqa ma'lumotlari telefon", "qarindoshlarning aloqa malumotlari telefon",
        ],
    },
    "parent_email": {
        "label": "Relative email",
        "aliases": [
            "parent email", "guardian email", "relative email", "электронная почта родителя",
            "qarindoshlarning aloqa ma'lumotlari elektron pochta", "qarindoshlarning aloqa malumotlari elektron pochta",
        ],
    },
}


def normalize_text(value: Any) -> str:
    text = unicodedata.normalize("NFKD", str(value or "")).lower()
    text = text.replace("ё", "е")
    text = re.sub(r"[ʻʼ’`´]", "'", text)
    text = re.sub(r"[^\w\s']+", " ", text, flags=re.UNICODE)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def compact_text(value: Any) -> str:
    return re.sub(r"[\s']+", "", normalize_text(value))


ALIAS_INDEX: list[tuple[str, str, str, str]] = []
for canonical_key, definition in FIELD_DEFINITIONS.items():
    for alias in [definition["label"], *definition["aliases"]]:
        ALIAS_INDEX.append((canonical_key, definition["label"], normalize_text(alias), compact_text(alias)))



def recognize_header(header: Any) -> ColumnInfo:
    original = str(header or "").strip()
    normalized = normalize_text(original)
    compact = compact_text(original)
    if not normalized:
        return ColumnInfo("", "", original, 0.0, "empty")

    # Context-aware contact headings produced by two-level school tables.
    if any(token in normalized for token in ("telefon", "phone", "телефон")) and any(
        token in normalized for token in ("qarindosh", "parent", "guardian", "родител")
    ):
        return ColumnInfo("parent_phone", FIELD_DEFINITIONS["parent_phone"]["label"], original, 0.98, "alias")
    if any(token in normalized for token in ("elektron pochta", "email", "e mail", "почта")) and any(
        token in normalized for token in ("qarindosh", "parent", "guardian", "родител")
    ):
        return ColumnInfo("parent_email", FIELD_DEFINITIONS["parent_email"]["label"], original, 0.98, "alias")

    for key, label, alias_norm, alias_compact in ALIAS_INDEX:
        if normalized == alias_norm or compact == alias_compact:
            return ColumnInfo(key, label, original, 0.99, "alias")

    header_tokens = {token for token in normalized.split() if len(token) >= 2}
    best_token: tuple[float, str, str] | None = None
    for key, label, alias_norm, _ in ALIAS_INDEX:
        alias_tokens = {token for token in alias_norm.split() if len(token) >= 2}
        if not header_tokens or not alias_tokens:
            continue
        intersection = header_tokens & alias_tokens
        containment = len(intersection) / min(len(header_tokens), len(alias_tokens))
        coverage = len(intersection) / max(len(header_tokens), len(alias_tokens))
        score = containment * 0.7 + coverage * 0.3
        if containment >= 0.75 and (best_token is None or score > best_token[0]):
            best_token = (score, key, label)
    if best_token:
        return ColumnInfo(best_token[1], best_token[2], original, min(0.94, 0.84 + best_token[0] * 0.1), "token")

    # Fuzzy matching is intentionally stricter than v0.4. Short headings such
    # as T/N must never accidentally become "Student name".
    if len(compact) >= 4:
        best: tuple[float, str, str] | None = None
        for key, label, _, alias_compact in ALIAS_INDEX:
            if len(alias_compact) < 4:
                continue
            score = SequenceMatcher(None, compact, alias_compact).ratio()
            if best is None or score > best[0]:
                best = (score, key, label)
        if best and best[0] >= 0.82:
            return ColumnInfo(best[1], best[2], original, min(0.90, best[0]), "fuzzy")

    slug = compact or secrets.token_hex(3)
    return ColumnInfo(f"custom_{slug}", original or "Untitled column", original, 0.5, "custom")

def canonical_base_key(key: str) -> str:
    """Remove duplicate-column suffixes without collapsing question numbers.

    ``question_1`` and ``question_2`` are different semantic fields, while
    ``student_name_2`` is a duplicate physical column of ``student_name``.
    """
    question_match = re.fullmatch(r"(question_\d+)(?:_\d+)?", key)
    if question_match:
        return question_match.group(1)
    return re.sub(r"_\d+$", "", key)


def cell_has_value(value: Any) -> bool:
    return value is not None and str(value).strip() != ""


def count_non_empty(row: list[Any]) -> int:
    return sum(1 for value in row if cell_has_value(value))


def header_row_score(row: list[Any]) -> float:
    """Score a likely heading row without mistaking numeric data for headers.

    Monitoring sheets often contain a first student row full of values such as
    ``1, 2-A, 72, 65, 60``. Numeric headings are supported later, but numeric
    *data rows* must not outrank a real textual heading row.
    """
    values = [str(value).strip() for value in row if cell_has_value(value)]
    if len(values) < 2:
        return float("-inf")
    normalized_values = [normalize_text(value) for value in values]
    unique_ratio = len(set(normalized_values)) / len(values)
    text_ratio = sum(bool(re.search(r"[A-Za-zА-Яа-яЎўҚқҒғҲҳ]", value)) for value in values) / len(values)
    recognitions = [recognize_header(value) for value in values]
    semantic_recognized = sum(item.method not in {"custom", "empty", "question"} for item in recognitions)
    numeric_ratio = sum(bool(re.fullmatch(r"[-+]?\d+(?:[.,]\d+)?", value.strip())) for value in values) / len(values)
    # Textual, unique, semantically recognizable rows are favored. Numeric-heavy
    # rows are strongly penalized because they are usually the first data row.
    return (
        len(values) * 1.4
        + unique_ratio * 2
        + text_ratio * 5
        + semantic_recognized * 3
        - numeric_ratio * 9
    )


def detect_header_row(rows: list[list[Any]]) -> int:
    if not rows:
        raise ValueError("The sheet is empty.")
    search_limit = min(len(rows), 25)
    scored = [(header_row_score(rows[index]), index) for index in range(search_limit)]
    score, index = max(scored, key=lambda item: item[0])
    if score == float("-inf"):
        raise ValueError("Could not find a usable header row in the first 25 rows.")
    return index


def make_unique_columns(headers: list[Any]) -> list[ColumnInfo]:
    used: dict[str, int] = {}
    result: list[ColumnInfo] = []
    for index, header in enumerate(headers):
        recognition = recognize_header(header or f"Column {index + 1}")
        base_key = recognition.key or f"column_{index + 1}"
        count = used.get(base_key, 0) + 1
        used[base_key] = count
        key = base_key if count == 1 else f"{base_key}_{count}"
        label = recognition.label if count == 1 else f"{recognition.label} {count}"
        result.append(ColumnInfo(key, label, recognition.original, recognition.confidence, recognition.method))
    return result


def normalize_cell_value(value: Any) -> Any:
    if isinstance(value, (datetime, date)):
        return value.isoformat()
    return value



def trim_empty_columns(rows: list[list[Any]]) -> list[list[Any]]:
    """Remove columns that are empty across the complete table."""
    if not rows:
        return []
    width = max((len(row) for row in rows), default=0)
    active = [
        index
        for index in range(width)
        if any(index < len(row) and cell_has_value(row[index]) for row in rows)
    ]
    return [[row[index] if index < len(row) else "" for index in active] for row in rows]


def looks_like_sequence(value: Any) -> bool:
    text = str(value or "").strip()
    if isinstance(value, float) and value.is_integer():
        return True
    return bool(re.fullmatch(r"\d{1,4}(?:\.0+)?", text))


def looks_like_person_name(value: Any) -> bool:
    text = normalize_text(value)
    if not text or looks_like_sequence(value):
        return False
    words = [word for word in text.split() if re.search(r"[A-Za-zА-Яа-яЎўҚқҒғҲҳ]", word)]
    return len(words) >= 2 and len(text) >= 5


def is_headerless_roster(rows: list[list[Any]]) -> bool:
    """Detect simple Word rosters where the first row is already a student."""
    rows = trim_empty_columns(rows)
    if len(rows) < 2 or max((len(row) for row in rows), default=0) < 2:
        return False
    first = rows[0]
    if len(first) >= 2:
        first_recognized = [recognize_header(first[0]), recognize_header(first[1])]
        if any(item.key in {"row_number", "student_name"} and item.method != "custom" for item in first_recognized):
            return False
    sample = rows[: min(20, len(rows))]
    seq_ratio = sum(looks_like_sequence(row[0] if row else "") for row in sample) / len(sample)
    name_ratio = sum(looks_like_person_name(row[1] if len(row) > 1 else "") for row in sample) / len(sample)
    return seq_ratio >= 0.7 and name_ratio >= 0.7


CLASS_PATTERN = re.compile(
    r"(?<!\d)(1[01]|[1-9])\s*(?:[-–—_]|\s)?\s*[,.'\"`«»„“”]*\s*([A-Za-zА-Яа-яЁёЎўҚқҒғҲҳ])(?=\b|[^A-Za-zА-Яа-яЁёЎўҚқҒғҲҳ])"
)


def extract_class_values(text: Any) -> list[str]:
    source = str(text or "")
    values: list[str] = []
    for match in CLASS_PATTERN.finditer(source):
        grade = int(match.group(1))
        if not 1 <= grade <= 11:
            continue
        value = f"{grade}-{match.group(2).upper()}"
        if value not in values:
            values.append(value)
    return values


def infer_class(contexts: list[tuple[str, str, int]]) -> dict[str, Any]:
    candidates: list[dict[str, Any]] = []
    for text, source, priority in contexts:
        for value in extract_class_values(text):
            candidates.append({"value": value, "source": source, "priority": priority})
    if not candidates:
        return {"value": None, "source": None, "conflict": False, "candidates": []}
    candidates.sort(key=lambda item: item["priority"], reverse=True)
    chosen = candidates[0]
    distinct = {item["value"] for item in candidates}
    return {
        "value": chosen["value"],
        "source": chosen["source"],
        "conflict": len(distinct) > 1,
        "candidates": [{"value": item["value"], "source": item["source"]} for item in candidates],
    }


def apply_inferred_class(
    columns: list[ColumnInfo], records: list[dict[str, Any]], class_info: dict[str, Any]
) -> tuple[list[ColumnInfo], list[dict[str, Any]]]:
    inferred = class_info.get("value")
    if not inferred:
        return columns, records
    class_column = next((column for column in columns if canonical_base_key(column.key) == "class"), None)
    if class_column is None:
        class_column = ColumnInfo("class", "Class", "Inferred class", 0.97, "inferred")
        columns = [*columns, class_column]
    for record in records:
        if not cell_has_value(record.get(class_column.key, "")):
            record[class_column.key] = inferred
    return columns, records


def records_from_explicit_headers(
    headers: list[Any], data_rows: list[list[Any]], source_file: str, source_sheet: str, start_row: int = 1
) -> tuple[list[ColumnInfo], list[dict[str, Any]]]:
    combined = trim_empty_columns([headers, *data_rows])
    if not combined:
        return [], []
    headers = combined[0]
    data_rows = combined[1:]
    columns = make_unique_columns(headers)
    records: list[dict[str, Any]] = []
    for offset, row in enumerate(data_rows, start=start_row):
        if count_non_empty(row) == 0 or is_repeated_header_row(row, headers):
            continue
        record: dict[str, Any] = {
            "__source_file": source_file,
            "__source_sheet": source_sheet,
            "__source_row": offset,
        }
        meaningful = 0
        for index, column in enumerate(columns):
            value = normalize_cell_value(row[index] if index < len(row) else "")
            record[column.key] = value
            meaningful += int(cell_has_value(value))
        if meaningful:
            records.append(record)
        if len(records) >= MAX_ROWS_PER_FILE:
            break
    return columns, records


def combine_values(left: Any, right: Any) -> Any:
    if not cell_has_value(left):
        return right
    if not cell_has_value(right) or str(left).strip() == str(right).strip():
        return left
    left_parts = [part.strip() for part in str(left).split(" | ")]
    if str(right).strip() in left_parts:
        return left
    return f"{left} | {right}"


def merge_continuation_records(records: list[dict[str, Any]], columns: list[ColumnInfo]) -> tuple[list[dict[str, Any]], int]:
    row_key = next((column.key for column in columns if canonical_base_key(column.key) == "row_number"), None)
    name_key = next((column.key for column in columns if canonical_base_key(column.key) == "student_name"), None)
    class_key = next((column.key for column in columns if canonical_base_key(column.key) == "class"), None)
    if not name_key:
        return records, 0
    merged: list[dict[str, Any]] = []
    merged_count = 0
    for record in records:
        identity = (
            str(record.get(row_key, "")).strip() if row_key else "",
            normalize_text(record.get(name_key, "")),
            normalize_text(record.get(class_key, "")) if class_key else "",
        )
        if merged:
            previous = merged[-1]
            previous_identity = (
                str(previous.get(row_key, "")).strip() if row_key else "",
                normalize_text(previous.get(name_key, "")),
                normalize_text(previous.get(class_key, "")) if class_key else "",
            )
            if identity[1] and identity == previous_identity:
                for column in columns:
                    previous[column.key] = combine_values(previous.get(column.key, ""), record.get(column.key, ""))
                merged_count += 1
                continue
        merged.append(record)
    return merged, merged_count

def is_repeated_header_row(row: list[Any], primary_header: list[Any]) -> bool:
    """Return True when a data-area row is actually another table header.

    School exports often repeat one- or two-line headings before every class
    section. For example, a secondary heading may contain only "Telefon" and
    "Elektron pochta". Without this check those labels become fake students.

    The rule is deliberately conservative: at least two non-empty cells must
    look like known field names, or at least two cells must repeat the primary
    header. Normal student rows such as "Ali Karimov | 5-E" do not match.
    """
    non_empty = [value for value in row if cell_has_value(value)]
    if len(non_empty) < 2:
        return False

    primary_values = {
        normalize_text(value)
        for value in primary_header
        if cell_has_value(value) and not looks_like_sequence(value)
    }
    exact_repeats = sum(
        1
        for value in non_empty
        if not looks_like_sequence(value) and normalize_text(value) in primary_values
    )

    recognized = [recognize_header(value) for value in non_empty]
    # Numeric monitoring answers such as 0/1 are data, even though numeric
    # question headers use the same visible values. They must not turn every
    # student row into a "repeated header".
    recognized_count = sum(
        info.method not in {"custom", "empty", "question"} for info in recognized
    )
    distinct_known_fields = {
        info.key for info in recognized if info.method not in {"custom", "empty", "question"}
    }

    if exact_repeats >= 2 and exact_repeats / len(non_empty) >= 0.5:
        return True

    if (
        recognized_count >= 2
        and len(distinct_known_fields) >= 2
        and recognized_count / len(non_empty) >= 0.67
    ):
        return True

    return False



def rows_to_records(rows: list[list[Any]], source_file: str, source_sheet: str) -> tuple[int, list[ColumnInfo], list[dict[str, Any]]]:
    rows = trim_empty_columns(rows)
    header_index = detect_header_row(rows)
    header_row = rows[header_index]
    columns, records = records_from_explicit_headers(
        header_row,
        rows[header_index + 1 :],
        source_file,
        source_sheet,
        start_row=header_index + 2,
    )
    return header_index, columns, records


def parse_xlsx(content: bytes, filename: str) -> dict[str, Any]:
    workbook = load_workbook(io.BytesIO(content), data_only=True, read_only=True)
    sheet = next((ws for ws in workbook.worksheets if ws.max_row and ws.max_column), None)
    if sheet is None:
        raise ValueError("Workbook contains no usable worksheet.")
    rows = [list(row) for row in sheet.iter_rows(values_only=True)]
    header_index, columns, records = rows_to_records(rows, filename, sheet.title)
    preheader = " ".join(str(value) for row in rows[:header_index] for value in row if cell_has_value(value))
    class_info = infer_class([(preheader, "document title", 100), (filename, "file name", 80), (sheet.title, "sheet name", 40)])
    columns, records = apply_inferred_class(columns, records, class_info)
    return {"sheet": sheet.title, "header_index": header_index, "columns": columns, "records": records, "class_info": class_info}


def parse_xls(content: bytes, filename: str) -> dict[str, Any]:
    if xlrd is None:
        raise ValueError("Old .xls support requires xlrd. Run: pip install -r requirements.txt")
    workbook = xlrd.open_workbook(file_contents=content)
    if workbook.nsheets == 0:
        raise ValueError("Workbook contains no worksheet.")
    sheet = next((workbook.sheet_by_index(i) for i in range(workbook.nsheets) if workbook.sheet_by_index(i).nrows), None)
    if sheet is None:
        raise ValueError("Workbook contains no usable worksheet.")
    rows = [sheet.row_values(i) for i in range(sheet.nrows)]
    header_index, columns, records = rows_to_records(rows, filename, sheet.name)
    preheader = " ".join(str(value) for row in rows[:header_index] for value in row if cell_has_value(value))
    class_info = infer_class([(preheader, "document title", 100), (filename, "file name", 80), (sheet.name, "sheet name", 40)])
    columns, records = apply_inferred_class(columns, records, class_info)
    return {"sheet": sheet.name, "header_index": header_index, "columns": columns, "records": records, "class_info": class_info}

def decode_csv(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1251", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")



def parse_csv(content: bytes, filename: str) -> dict[str, Any]:
    text = decode_csv(content)
    sample = text[:4096]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
        delimiter = dialect.delimiter
    except csv.Error:
        delimiter = ","
    rows = [row for row in csv.reader(io.StringIO(text), delimiter=delimiter) if any(cell.strip() for cell in row)]
    header_index, columns, records = rows_to_records(rows, filename, "CSV")
    class_info = infer_class([(filename, "file name", 80)])
    columns, records = apply_inferred_class(columns, records, class_info)
    return {"sheet": "CSV", "header_index": header_index, "columns": columns, "records": records, "class_info": class_info}

OLE2_MAGIC = bytes.fromhex("D0CF11E0A1B11AE1")
ZIP_MAGIC = b"PK\x03\x04"


def sniff_zip_office_type(content: bytes) -> str:
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            names = set(archive.namelist())
    except (zipfile.BadZipFile, OSError):
        return "zip"
    if "word/document.xml" in names:
        return "docx-zip"
    if "xl/workbook.xml" in names:
        return "xlsx-zip"
    return "zip"


def sniff_file_type(content: bytes, filename: str) -> str:
    if content.startswith(OLE2_MAGIC):
        return "ole2"
    if content.startswith(ZIP_MAGIC):
        return sniff_zip_office_type(content)

    stripped = content.lstrip()
    lowered = stripped[:4096].lower()
    if lowered.startswith(b"<?xml") or lowered.startswith(b"<workbook"):
        if b"urn:schemas-microsoft-com:office:spreadsheet" in lowered or b"<workbook" in lowered:
            return "spreadsheetml-xml"
        return "xml"
    if b"<html" in lowered or b"<table" in lowered:
        return "html-table"

    try:
        text = decode_csv(content)
        first_lines = [line for line in text.splitlines() if line.strip()][:5]
        if first_lines and any(delim in first_lines[0] for delim in ("\t", ";", ",", "|")):
            return "delimited-text"
    except Exception:
        pass

    return "unknown"



def parse_html_table(content: bytes, filename: str) -> dict[str, Any]:
    text = decode_csv(content)
    soup = BeautifulSoup(text, "html.parser")
    tables = soup.find_all("table")
    if not tables:
        raise ValueError("The file looks like HTML, but no table was found.")

    def table_size(table: Any) -> int:
        return sum(len(row.find_all(["th", "td"], recursive=False)) for row in table.find_all("tr"))

    table = max(tables, key=table_size)
    source_rows = table.find_all("tr")
    grid: dict[tuple[int, int], str] = {}
    row_types: list[str] = []
    max_column = 0

    for row_index, tr in enumerate(source_rows):
        cells = tr.find_all(["th", "td"], recursive=False)
        row_types.append("header" if cells and all(cell.name == "th" for cell in cells) else "data")
        column_index = 0
        for cell in cells:
            while (row_index, column_index) in grid:
                column_index += 1
            value = unescape(cell.get_text(" ", strip=True))
            rowspan = max(1, int(cell.get("rowspan", 1) or 1))
            colspan = max(1, int(cell.get("colspan", 1) or 1))
            for target_row in range(row_index, row_index + rowspan):
                for target_column in range(column_index, column_index + colspan):
                    grid[(target_row, target_column)] = value
                    max_column = max(max_column, target_column + 1)
            column_index += colspan

    matrix = [
        [grid.get((row_index, column_index), "") for column_index in range(max_column)]
        for row_index in range(len(source_rows))
    ]
    matrix = trim_empty_columns(matrix)
    if not matrix:
        raise ValueError("The HTML table is empty.")

    first_header = next((index for index, kind in enumerate(row_types) if kind == "header"), None)
    if first_header is None:
        header_index, columns, records = rows_to_records(matrix, filename, "HTML table")
        header_mode = "automatic"
        header_rows = 1
    else:
        last_header = first_header
        while last_header + 1 < len(row_types) and row_types[last_header + 1] == "header":
            last_header += 1
        header_rows = last_header - first_header + 1
        headers: list[str] = []
        for column_index in range(len(matrix[0])):
            parts: list[str] = []
            for row_index in range(first_header, last_header + 1):
                value = str(matrix[row_index][column_index] or "").strip()
                if value and (not parts or normalize_text(value) != normalize_text(parts[-1])):
                    parts.append(value)
            headers.append(" / ".join(parts) if parts else f"Column {column_index + 1}")
        columns, records = records_from_explicit_headers(
            headers,
            matrix[last_header + 1 :],
            filename,
            "HTML table",
            start_row=last_header + 2,
        )
        header_index = last_header
        header_mode = "multi_row" if header_rows > 1 else "explicit"

    class_info = infer_class([(filename, "file name", 80)])
    columns, records = apply_inferred_class(columns, records, class_info)
    records, merged_count = merge_continuation_records(records, columns)
    return {
        "sheet": "HTML table",
        "header_index": header_index,
        "columns": columns,
        "records": records,
        "detected_type": "HTML table saved as .xls",
        "class_info": class_info,
        "header_mode": header_mode,
        "header_rows": header_rows,
        "continuation_rows_merged": merged_count,
    }


def parse_spreadsheetml(content: bytes, filename: str) -> dict[str, Any]:
    try:
        root = ET.fromstring(content)
    except ET.ParseError as exc:
        raise ValueError(f"Invalid SpreadsheetML XML: {exc}") from exc

    ns = {"ss": "urn:schemas-microsoft-com:office:spreadsheet"}
    worksheets = root.findall(".//ss:Worksheet", ns)
    if not worksheets:
        worksheets = [node for node in root.iter() if node.tag.endswith("Worksheet")]
    if not worksheets:
        raise ValueError("SpreadsheetML workbook contains no worksheet.")

    worksheet = worksheets[0]
    sheet_name = worksheet.attrib.get("{urn:schemas-microsoft-com:office:spreadsheet}Name") or worksheet.attrib.get("Name") or "SpreadsheetML"
    row_nodes = worksheet.findall(".//ss:Row", ns) or [node for node in worksheet.iter() if node.tag.endswith("Row")]
    rows: list[list[Any]] = []
    for row_node in row_nodes:
        row: list[Any] = []
        cell_nodes = list(row_node.findall("ss:Cell", ns)) or [node for node in row_node if node.tag.endswith("Cell")]
        for cell in cell_nodes:
            index_value = cell.attrib.get("{urn:schemas-microsoft-com:office:spreadsheet}Index") or cell.attrib.get("Index")
            if index_value:
                while len(row) < int(index_value) - 1:
                    row.append("")
            data_node = cell.find("ss:Data", ns)
            if data_node is None:
                data_node = next((node for node in cell.iter() if node.tag.endswith("Data")), None)
            row.append(data_node.text if data_node is not None and data_node.text is not None else "")
        if row:
            rows.append(row)

    header_index, columns, records = rows_to_records(rows, filename, sheet_name)
    preheader = " ".join(str(value) for row in rows[:header_index] for value in row if cell_has_value(value))
    class_info = infer_class([(preheader, "document title", 100), (filename, "file name", 80), (sheet_name, "sheet name", 40)])
    columns, records = apply_inferred_class(columns, records, class_info)
    return {
        "sheet": sheet_name,
        "header_index": header_index,
        "columns": columns,
        "records": records,
        "detected_type": "Excel 2003 XML / SpreadsheetML",
        "class_info": class_info,
    }

def parse_delimited_disguised_as_xls(content: bytes, filename: str) -> dict[str, Any]:
    parsed = parse_csv(content, filename)
    parsed["detected_type"] = "Delimited text saved with .xls extension"
    return parsed


def word_table_to_rows(table: Any) -> list[list[Any]]:
    rows: list[list[Any]] = []
    for table_row in table.rows:
        values = [re.sub(r"\s+", " ", cell.text or "").strip() for cell in table_row.cells]
        if any(cell_has_value(value) for value in values):
            rows.append(values)
    return rows


def table_signature(columns: list[ColumnInfo]) -> tuple[str, ...]:
    signature: list[str] = []
    for column in columns:
        base_key = canonical_base_key(column.key)
        if base_key.startswith("custom_"):
            base_key = f"custom:{normalize_text(column.original)}"
        signature.append(base_key)
    return tuple(signature)


def table_quality(columns: list[ColumnInfo], records: list[dict[str, Any]]) -> float:
    recognized = sum(column.method not in {"custom", "empty"} for column in columns)
    return len(records) * max(1, len(columns)) + recognized * 12



def parse_docx(content: bytes, filename: str) -> dict[str, Any]:
    try:
        document = Document(io.BytesIO(content))
    except Exception as exc:
        raise ValueError(f"Could not open the DOCX document: {exc}") from exc
    if not document.tables:
        raise ValueError("The Word document contains no tables.")

    document_text = " ".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
    class_info = infer_class([(document_text, "document title", 100), (filename, "file name", 80)])
    valid_tables: list[dict[str, Any]] = []
    table_summaries: list[dict[str, Any]] = []

    for table_index, table in enumerate(document.tables, start=1):
        rows = trim_empty_columns(word_table_to_rows(table))
        summary: dict[str, Any] = {
            "table": table_index,
            "raw_rows": len(rows),
            "raw_columns": max((len(row) for row in rows), default=0),
        }
        if not rows:
            summary.update({"usable": False, "error": "Empty table"})
            table_summaries.append(summary)
            continue
        try:
            if is_headerless_roster(rows):
                headers = ["№", "O‘quvchi F.I.Sh."] + [f"Column {index + 1}" for index in range(2, len(rows[0]))]
                columns, records = records_from_explicit_headers(headers, rows, filename, f"Word table {table_index}", start_row=1)
                header_index = -1
                header_mode = "inferred_roster"
            else:
                header_index, columns, records = rows_to_records(rows, filename, f"Word table {table_index}")
                header_mode = "explicit"
            columns, records = apply_inferred_class(columns, records, class_info)
            records, merged_count = merge_continuation_records(records, columns)
            signature = table_signature(columns)
            summary.update({
                "usable": True,
                "header_row": header_index + 1 if header_index >= 0 else None,
                "header_mode": header_mode,
                "rows": len(records),
                "columns": len(columns),
                "headers": [column.original for column in columns],
                "signature": list(signature),
                "continuation_rows_merged": merged_count,
            })
            valid_tables.append({
                "index": table_index,
                "header_index": header_index,
                "header_mode": header_mode,
                "columns": columns,
                "records": records,
                "signature": signature,
                "quality": table_quality(columns, records),
            })
        except Exception as exc:
            summary.update({"usable": False, "error": str(exc)})
        table_summaries.append(summary)

    if not valid_tables:
        raise ValueError(f"The document contains {len(document.tables)} table(s), but none looked like a usable data table.")

    signatures = {entry["signature"] for entry in valid_tables}
    tables_differ = len(signatures) > 1
    tables_merged = len(valid_tables) > 1 and not tables_differ
    if tables_merged:
        selected = valid_tables[0]
        records: list[dict[str, Any]] = []
        for entry in valid_tables:
            records.extend(entry["records"])
        selected_table = None
        sheet = f"{len(valid_tables)} similar Word tables"
    else:
        selected = max(valid_tables, key=lambda entry: entry["quality"])
        records = selected["records"]
        selected_table = selected["index"]
        sheet = f"Word table {selected_table}"

    warning_code: str | None = None
    if tables_differ:
        warning_code = "different_word_tables"
    elif tables_merged:
        warning_code = "similar_word_tables_merged"

    return {
        "sheet": sheet,
        "header_index": selected["header_index"],
        "header_mode": selected.get("header_mode", "explicit"),
        "columns": selected["columns"],
        "records": records,
        "detected_type": "DOCX Word document",
        "source_kind": "word",
        "table_count": len(document.tables),
        "usable_table_count": len(valid_tables),
        "tables_differ": tables_differ,
        "tables_merged": tables_merged,
        "selected_table": selected_table,
        "warning_code": warning_code,
        "table_summaries": table_summaries,
        "class_info": class_info,
    }

def find_libreoffice() -> str | None:
    candidates = [
        shutil.which("soffice"),
        shutil.which("libreoffice"),
        r"C:\\Program Files\\LibreOffice\\program\\soffice.exe",
        r"C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return str(candidate)
    return None


def parse_doc(content: bytes, filename: str) -> dict[str, Any]:
    converter = find_libreoffice()
    if converter is None:
        raise ValueError(
            "Legacy .doc import requires LibreOffice. Install LibreOffice, restart FormaFlow, "
            "or open the file in Word and save it as .docx."
        )

    with tempfile.TemporaryDirectory(prefix="forma_flow_doc_") as temp_dir:
        temp_path = Path(temp_dir)
        source_path = temp_path / safe_filename(filename)
        if source_path.suffix.lower() != ".doc":
            source_path = source_path.with_suffix(".doc")
        source_path.write_bytes(content)

        command = [
            converter,
            "--headless",
            "--convert-to",
            "docx",
            "--outdir",
            str(temp_path),
            str(source_path),
        ]
        try:
            result = subprocess.run(
                command,
                capture_output=True,
                text=True,
                timeout=60,
                check=False,
            )
        except (OSError, subprocess.TimeoutExpired) as exc:
            raise ValueError(f"LibreOffice could not convert the .doc file: {exc}") from exc

        converted = source_path.with_suffix(".docx")
        if not converted.exists():
            candidates = list(temp_path.glob("*.docx"))
            converted = candidates[0] if candidates else converted
        if not converted.exists():
            details = (result.stderr or result.stdout or "No converter output").strip()
            raise ValueError(f"LibreOffice could not convert this .doc file. {details}")

        parsed = parse_docx(converted.read_bytes(), filename)
        parsed["detected_type"] = "Legacy DOC converted locally with LibreOffice"
        return parsed


def parse_uploaded_file(filename: str, content: bytes) -> dict[str, Any]:
    if not content:
        raise ValueError("The selected file is empty (0 bytes).")

    suffix = Path(filename).suffix.lower()
    detected = sniff_file_type(content, filename)

    if detected == "docx-zip":
        return parse_docx(content, filename)
    if detected == "xlsx-zip":
        parsed = parse_xlsx(content, filename)
        parsed["detected_type"] = "XLSX/ZIP workbook" + (" renamed" if suffix not in {".xlsx", ".xlsm"} else "")
        parsed["source_kind"] = "spreadsheet"
        return parsed
    if detected == "ole2":
        if suffix == ".doc":
            return parse_doc(content, filename)
        parsed = parse_xls(content, filename)
        parsed["detected_type"] = "Legacy binary XLS workbook"
        parsed["source_kind"] = "spreadsheet"
        return parsed
    if detected == "html-table":
        parsed = parse_html_table(content, filename)
        parsed["source_kind"] = "spreadsheet"
        return parsed
    if detected == "spreadsheetml-xml":
        parsed = parse_spreadsheetml(content, filename)
        parsed["source_kind"] = "spreadsheet"
        return parsed
    if detected == "delimited-text":
        parsed = parse_delimited_disguised_as_xls(content, filename)
        parsed["source_kind"] = "spreadsheet"
        return parsed

    try:
        if suffix == ".docx":
            return parse_docx(content, filename)
        if suffix == ".doc":
            return parse_doc(content, filename)
        if suffix in {".xlsx", ".xlsm"}:
            parsed = parse_xlsx(content, filename)
            parsed["detected_type"] = "XLSX workbook"
            parsed["source_kind"] = "spreadsheet"
            return parsed
        if suffix == ".xls":
            parsed = parse_xls(content, filename)
            parsed["detected_type"] = "Legacy binary XLS workbook"
            parsed["source_kind"] = "spreadsheet"
            return parsed
        if suffix == ".csv":
            parsed = parse_csv(content, filename)
            parsed["detected_type"] = "CSV text"
            parsed["source_kind"] = "spreadsheet"
            return parsed
    except Exception as exc:
        prefix = content[:12].hex(" ").upper()
        raise ValueError(
            f"Could not read this file as {suffix or 'a supported document'}. "
            f"Detected content: {detected}. First bytes: {prefix}. Original error: {exc}"
        ) from exc

    raise ValueError(
        f"Unsupported file. Extension: {suffix or 'none'}; detected content: {detected}. "
        "Supported inputs are XLSX, XLS, XLSM, CSV, DOCX, and DOC (DOC requires LibreOffice)."
    )



def merge_parsed_files(parsed_files: list[dict[str, Any]]) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    column_map: OrderedDict[str, dict[str, Any]] = OrderedDict()
    records: list[dict[str, Any]] = []
    core_default = {"row_number", "student_name", "class", "grade"}
    sensitive = {"pinfl", "birth_date", "phone", "email", "address", "parent_phone", "parent_email"}

    for parsed in parsed_files:
        for column in parsed["columns"]:
            base_key = canonical_base_key(column.key)
            if column.key not in column_map:
                column_map[column.key] = {
                    "key": column.key,
                    "label": column.label,
                    "originals": [column.original],
                    "confidence": column.confidence,
                    "method": column.method,
                    "default_selected": base_key in core_default,
                    "sensitive": base_key in sensitive,
                }
            else:
                existing = column_map[column.key]
                if column.original not in existing["originals"]:
                    existing["originals"].append(column.original)
                existing["confidence"] = max(existing["confidence"], column.confidence)
        records.extend(parsed["records"])

    columns = list(column_map.values())
    if columns and not any(column["default_selected"] for column in columns):
        for column in columns[: min(3, len(columns))]:
            column["default_selected"] = True
    return columns, records

class FixedColumn(BaseModel):
    name: str = Field(min_length=1, max_length=100)
    value: str = Field(default="", max_length=500)


class OutputColumn(BaseModel):
    key: str
    name: str = Field(min_length=1, max_length=100)


class DerivedColumn(BaseModel):
    name: str = Field(min_length=1, max_length=100)
    kind: Literal["next_class", "sequence", "add_number", "copy"]
    source_key: str | None = None
    stop_grades: list[int] = Field(default_factory=lambda: [9, 11])
    amount: float = 1
    start: int = 1


class ProcessingOptions(BaseModel):
    trim_whitespace: bool = True
    remove_duplicates: bool = False
    skip_blank_key: str | None = None
    sort_key: str | None = None


class BuildRequest(BaseModel):
    session_id: str
    columns: list[OutputColumn]
    fixed_columns: list[FixedColumn] = Field(default_factory=list)
    derived_columns: list[DerivedColumn] = Field(default_factory=list)
    options: ProcessingOptions = Field(default_factory=ProcessingOptions)


class PreviewRequest(BuildRequest):
    limit: int = Field(default=50, ge=1, le=200)


class ExportRequest(BuildRequest):
    format: str


def clean_output_value(value: Any, trim_whitespace: bool) -> Any:
    if not trim_whitespace or not isinstance(value, str):
        return value
    return re.sub(r"\s+", " ", value).strip()


def next_class_value(value: Any, stop_grades: list[int] | None = None) -> str:
    """Return the following class while preserving the class suffix.

    Examples: 5-b -> 6-b, 8-A -> 9-A, 10 Б -> 11 Б.
    Terminal grades (9 and 11 by default) produce an empty value.
    """
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        text = str(int(value))
    else:
        text = str(value)
    match = re.match(r"^\s*(\d{1,2})(.*)$", text)
    if not match:
        return ""
    grade = int(match.group(1))
    suffix = match.group(2)
    terminal = set(stop_grades or [9, 11])
    if grade in terminal or grade < 1 or grade >= 11:
        return ""
    return f"{grade + 1}{suffix}"


def add_number_value(value: Any, amount: float) -> Any:
    if value is None or str(value).strip() == "":
        return ""
    try:
        number = float(str(value).replace(",", ".").strip())
    except (TypeError, ValueError):
        return ""
    result = number + amount
    return int(result) if result.is_integer() else result


def natural_sort_value(value: Any) -> tuple[Any, ...]:
    text = str(value or "").strip().lower()
    parts = re.split(r"(\d+(?:[.,]\d+)?)", text)
    result: list[Any] = []
    for part in parts:
        if not part:
            continue
        try:
            result.append((0, float(part.replace(",", "."))))
        except ValueError:
            result.append((1, part))
    return tuple(result)


def derive_value(record: dict[str, Any], derived: DerivedColumn, row_number: int) -> Any:
    source = record.get(derived.source_key or "", "")
    if derived.kind == "next_class":
        return next_class_value(source, derived.stop_grades)
    if derived.kind == "sequence":
        return derived.start + row_number
    if derived.kind == "add_number":
        return add_number_value(source, derived.amount)
    if derived.kind == "copy":
        return source
    return ""


def build_output_rows(session: dict[str, Any], request: BuildRequest) -> list[dict[str, Any]]:
    available = {column["key"] for column in session["columns"]}
    requested = [column for column in request.columns if column.key in available]
    if not requested:
        raise HTTPException(status_code=400, detail="Select at least one source column.")

    records = list(session["records"])
    options = request.options

    if options.skip_blank_key and options.skip_blank_key in available:
        records = [record for record in records if cell_has_value(record.get(options.skip_blank_key, ""))]

    if options.sort_key and options.sort_key in available:
        records.sort(key=lambda record: natural_sort_value(record.get(options.sort_key, "")))

    output_rows: list[dict[str, Any]] = []
    seen: set[tuple[str, ...]] = set()

    for row_number, record in enumerate(records):
        output: OrderedDict[str, Any] = OrderedDict()
        for column in requested:
            output[column.name] = clean_output_value(
                record.get(column.key, ""), options.trim_whitespace
            )
        for fixed in request.fixed_columns:
            output[fixed.name] = clean_output_value(fixed.value, options.trim_whitespace)
        for derived in request.derived_columns:
            if derived.kind != "sequence" and derived.source_key not in available:
                continue
            output[derived.name] = clean_output_value(
                derive_value(record, derived, row_number), options.trim_whitespace
            )

        output_dict = dict(output)
        if options.remove_duplicates:
            fingerprint = tuple(str(value) for value in output_dict.values())
            if fingerprint in seen:
                continue
            seen.add(fingerprint)
        output_rows.append(output_dict)

    return output_rows


def safe_filename(name: str) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', "_", name).strip(" ._")
    return cleaned[:100] or "FormaFlow_Output"


def create_xlsx(rows: list[dict[str, Any]]) -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "FormaFlow"
    headers = list(rows[0].keys())
    sheet.append(headers)
    header_fill = PatternFill("solid", fgColor="635BFF")
    for cell in sheet[1]:
        cell.font = Font(color="FFFFFF", bold=True)
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")
    for row in rows:
        sheet.append([row.get(header, "") for header in headers])
    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = sheet.dimensions
    for index, header in enumerate(headers, start=1):
        max_length = len(str(header))
        for cell in sheet[get_column_letter(index)][1:251]:
            max_length = max(max_length, len(str(cell.value or "")))
        sheet.column_dimensions[get_column_letter(index)].width = min(40, max(10, max_length + 2))
    output = io.BytesIO()
    workbook.save(output)
    return output.getvalue()


def create_csv(rows: list[dict[str, Any]]) -> bytes:
    output = io.StringIO(newline="")
    writer = csv.DictWriter(output, fieldnames=list(rows[0].keys()))
    writer.writeheader()
    writer.writerows(rows)
    return ("\ufeff" + output.getvalue()).encode("utf-8")


def create_docx(rows: list[dict[str, Any]]) -> bytes:
    document = Document()
    section = document.sections[0]
    headers = list(rows[0].keys())
    if len(headers) > 5:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("FormaFlow Export")
    run.bold = True
    run.font.size = Pt(16)

    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for index, header in enumerate(headers):
            cells[index].text = str(row.get(header, "") or "")
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


# ---------------------------------------------------------------------------
# FormaFlow v0.6 - Document Intelligence engine
# ---------------------------------------------------------------------------

import hashlib
from collections import defaultdict

try:
    import pdfplumber
except ImportError:  # PDF support is optional until requirements are installed.
    pdfplumber = None

app.version = "0.6.0"

# Extend the canonical schema. These aliases are deliberately deterministic:
# the application never sends school data to an external AI service.
FIELD_DEFINITIONS.update({
    "promoted_class": {
        "label": "Promoted / next class",
        "aliases": [
            "promoted class", "next class", "new class", "ko'chirilgan sinfi",
            "ko‘chirilgan sinfi", "kochirilgan sinfi", "keyingi sinf",
            "переведен в класс", "следующий класс",
        ],
    },
    "total": {
        "label": "Total",
        "aliases": ["total", "jami", "umumiy bali", "umumiy ball", "итого", "сумма"],
    },
    "percentage": {
        "label": "Percentage",
        "aliases": ["percentage", "percent", "foiz", "фоиз", "процент", "%"],
    },
    "subject": {
        "label": "Subject",
        "aliases": ["subject", "fan", "предмет"],
    },
    "teacher": {
        "label": "Teacher",
        "aliases": ["teacher", "fan o'qituvchisi", "fan o‘qituvchisi", "o'qituvchi", "учитель", "преподаватель"],
    },
    "place": {
        "label": "Place",
        "aliases": ["place", "rank", "o'rni", "o‘rni", "орын", "место"],
    },
    "score": {
        "label": "Score",
        "aliases": ["score", "ball", "bali", "olgan bali", "natija", "result", "балл", "результат"],
    },
    "academic_year": {
        "label": "Academic year",
        "aliases": ["academic year", "school year", "o'quv yili", "o‘quv yili", "учебный год"],
    },
})

FIELD_DEFINITIONS["row_number"]["aliases"].extend(["t/r", "t r", "tr", "t\\r"])
FIELD_DEFINITIONS["student_name"]["aliases"].extend([
    "ism sharifi", "ism familiyasi", "ism va familiyasi", "ism familyasi",
    "o'quvchilar ismi va familyasi", "o‘quvchilar ismi va familyasi",
    "o'quvchining ism va familiyasi", "o‘quvchining ism va familiyasi",
    "o'quvchining ismi familiyasi", "o‘quvchining ismi familiyasi",
    "o'quvchilarning ismi familiyasi otasining ismi",
    "исм фамилияси", "исм шарифи",
])


def rebuild_alias_index() -> None:
    global ALIAS_INDEX
    ALIAS_INDEX = []
    for canonical_key, definition in FIELD_DEFINITIONS.items():
        for alias in [definition["label"], *definition.get("aliases", [])]:
            ALIAS_INDEX.append(
                (canonical_key, definition["label"], normalize_text(alias), compact_text(alias))
            )


rebuild_alias_index()


def recognize_header(header: Any) -> ColumnInfo:
    original = str(header or "").strip()
    normalized = normalize_text(original)
    compact = compact_text(original)
    if not normalized and original.strip() != "%":
        return ColumnInfo("", "", original, 0.0, "empty")

    # Numeric monitoring-question headings are a real schema, not random custom fields.
    numeric_match = re.fullmatch(r"\s*(\d{1,3})\s*", original)
    if numeric_match:
        number = numeric_match.group(1)
        return ColumnInfo(f"question_{number}", number, original, 0.99, "question")

    if original.strip() == "%" or normalized in {"foiz", "фоиз", "percent", "percentage", "процент"}:
        return ColumnInfo("percentage", FIELD_DEFINITIONS["percentage"]["label"], original, 0.99, "alias")

    # High-value phrases that token matching can otherwise misunderstand.
    phrase_rules = [
        ("promoted_class", ("ko'chirilgan sinfi", "ko chirilgan sinfi", "keyingi sinf", "promoted class", "next class")),
        ("student_name", ("ism sharifi", "ism familiyasi", "ism familyasi", "f i sh", "фио", "исм фамилияси", "исм шарифи")),
        ("row_number", ("t r", "t n", "tartib raqam", "п п")),
    ]
    for key, phrases in phrase_rules:
        if any(phrase in normalized for phrase in phrases):
            return ColumnInfo(key, FIELD_DEFINITIONS[key]["label"], original, 0.98, "alias")

    # Context-aware guardian contact headings produced by two-level tables.
    if any(token in normalized for token in ("telefon", "phone", "телефон")) and any(
        token in normalized for token in ("qarindosh", "parent", "guardian", "родител")
    ):
        return ColumnInfo("parent_phone", FIELD_DEFINITIONS["parent_phone"]["label"], original, 0.98, "alias")
    if any(token in normalized for token in ("elektron pochta", "email", "e mail", "почта")) and any(
        token in normalized for token in ("qarindosh", "parent", "guardian", "родител")
    ):
        return ColumnInfo("parent_email", FIELD_DEFINITIONS["parent_email"]["label"], original, 0.98, "alias")

    for key, label, alias_norm, alias_compact in ALIAS_INDEX:
        if normalized == alias_norm or compact == alias_compact:
            return ColumnInfo(key, label, original, 0.99, "alias")

    header_tokens = {token for token in normalized.split() if len(token) >= 2}
    best_token: tuple[float, str, str] | None = None
    for key, label, alias_norm, _ in ALIAS_INDEX:
        alias_tokens = {token for token in alias_norm.split() if len(token) >= 2}
        if not header_tokens or not alias_tokens:
            continue
        intersection = header_tokens & alias_tokens
        containment = len(intersection) / min(len(header_tokens), len(alias_tokens))
        coverage = len(intersection) / max(len(header_tokens), len(alias_tokens))
        score = containment * 0.7 + coverage * 0.3
        if containment >= 0.78 and (best_token is None or score > best_token[0]):
            best_token = (score, key, label)
    if best_token:
        return ColumnInfo(best_token[1], best_token[2], original, min(0.94, 0.84 + best_token[0] * 0.1), "token")

    if len(compact) >= 4:
        best: tuple[float, str, str] | None = None
        for key, label, _, alias_compact in ALIAS_INDEX:
            if len(alias_compact) < 4:
                continue
            score = SequenceMatcher(None, compact, alias_compact).ratio()
            if best is None or score > best[0]:
                best = (score, key, label)
        if best and best[0] >= 0.84:
            return ColumnInfo(best[1], best[2], original, min(0.90, best[0]), "fuzzy")

    slug = compact or secrets.token_hex(3)
    return ColumnInfo(f"custom_{slug}", original or "Untitled column", original, 0.5, "custom")


def collapse_duplicate_columns(rows: list[list[Any]]) -> tuple[list[list[Any]], int]:
    """Remove empty and duplicated logical columns created by Word merged cells.

    python-docx can expose one visually merged cell as several columns. Header
    rows may be blank in the duplicated copies, so equality is evaluated using
    non-conflicting values rather than requiring byte-for-byte equality.
    """
    if not rows:
        return [], 0
    width = max((len(row) for row in rows), default=0)
    padded = [[row[i] if i < len(row) else "" for i in range(width)] for row in rows]
    active = [i for i in range(width) if any(cell_has_value(row[i]) for row in padded)]
    kept: list[int] = []
    removed = 0

    def compatible_duplicate(left: int, right: int) -> bool:
        matching = 0
        conflicting = 0
        for row in padded:
            a = normalize_text(row[left])
            b = normalize_text(row[right])
            if a and b:
                if a == b:
                    matching += 1
                else:
                    conflicting += 1
        return conflicting == 0 and matching >= 2

    for index in active:
        duplicate_of = next((kept_index for kept_index in kept if compatible_duplicate(kept_index, index)), None)
        if duplicate_of is not None:
            # Preserve a non-empty value from either physical column.
            for row in padded:
                if not cell_has_value(row[duplicate_of]) and cell_has_value(row[index]):
                    row[duplicate_of] = row[index]
            removed += 1
            continue
        kept.append(index)
    return [[row[index] for index in kept] for row in padded], removed


def trim_empty_columns(rows: list[list[Any]]) -> list[list[Any]]:
    cleaned, _ = collapse_duplicate_columns(rows)
    return cleaned


def looks_like_class(value: Any) -> bool:
    return bool(extract_class_values(value))


def find_base_key(columns: list[ColumnInfo], base: str) -> str | None:
    return next((column.key for column in columns if canonical_base_key(column.key) == base), None)


def explicit_roster_header_index(rows: list[list[Any]]) -> int | None:
    """Find sparse roster headers such as ['', 'O‘quvchining ism va familiyasi']."""
    for index, row in enumerate(rows[:15]):
        recognized = [recognize_header(value) for value in row]
        known = [item for item in recognized if item.method not in {"custom", "empty"}]
        has_name = any(item.key == "student_name" for item in known)
        if not has_name:
            continue
        following = rows[index + 1 : index + 9]
        if not following:
            continue
        sequence_ratio = sum(looks_like_sequence(r[0] if r else "") for r in following) / len(following)
        person_ratio = max(
            (
                sum(looks_like_person_name(r[col] if col < len(r) else "") for r in following) / len(following)
                for col in range(max((len(r) for r in following), default=0))
            ),
            default=0,
        )
        if person_ratio >= 0.55 and (sequence_ratio >= 0.45 or len(row) == 1):
            return index
    return None


def detect_header_row(rows: list[list[Any]]) -> int:
    if not rows:
        raise ValueError("The table is empty.")
    sparse = explicit_roster_header_index(rows)
    if sparse is not None:
        return sparse
    search_limit = min(len(rows), 30)
    scored = [(header_row_score(rows[index]), index) for index in range(search_limit)]
    score, index = max(scored, key=lambda item: item[0])
    if score == float("-inf"):
        raise ValueError("Could not find a usable header row in the first 30 rows.")
    return index


def make_unique_columns(headers: list[Any]) -> list[ColumnInfo]:
    used: dict[str, int] = {}
    result: list[ColumnInfo] = []
    for index, header in enumerate(headers):
        recognition = recognize_header(header or f"Column {index + 1}")
        base_key = recognition.key or f"column_{index + 1}"
        count = used.get(base_key, 0) + 1
        used[base_key] = count
        key = base_key if count == 1 else f"{base_key}_{count}"
        label = recognition.label if count == 1 else f"{recognition.label} {count}"
        result.append(ColumnInfo(key, label, recognition.original, recognition.confidence, recognition.method))
    return result


def is_footer_or_label(value: Any) -> bool:
    text = normalize_text(value)
    if not text:
        return False
    footer_phrases = (
        "sinf rahbari", "o ibdo", "o'ibdo", "директор", "teacher", "o'qituvchi",
        "jami", "итого", "saldo", "сальдо", "imzo", "подпись", "kotib", "rais",
    )
    return any(phrase in text for phrase in footer_phrases)


def infer_columns_from_values(columns: list[ColumnInfo], records: list[dict[str, Any]]) -> list[ColumnInfo]:
    """Recover meanings for blank headers using the values beneath them."""
    occupied = {canonical_base_key(column.key) for column in columns if column.method != "custom"}
    updated: list[ColumnInfo] = []
    for column in columns:
        base = canonical_base_key(column.key)
        if column.method != "custom" and base not in {""}:
            updated.append(column)
            continue
        values = [record.get(column.key, "") for record in records[:100] if cell_has_value(record.get(column.key, ""))]
        if not values:
            updated.append(column)
            continue
        sequence_ratio = sum(looks_like_sequence(value) for value in values) / len(values)
        person_ratio = sum(looks_like_person_name(value) for value in values) / len(values)
        class_ratio = sum(looks_like_class(value) for value in values) / len(values)
        enough_evidence = len(values) >= 3
        replacement: ColumnInfo | None = None
        if enough_evidence and sequence_ratio >= 0.75 and "row_number" not in occupied:
            replacement = ColumnInfo("row_number", "No.", column.original or "Inferred row number", 0.9, "inferred")
        elif enough_evidence and person_ratio >= 0.75 and "student_name" not in occupied:
            replacement = ColumnInfo("student_name", "Student name", column.original or "Inferred student name", 0.88, "inferred")
        elif enough_evidence and class_ratio >= 0.75:
            target = "class" if "class" not in occupied else "promoted_class" if "promoted_class" not in occupied else None
            if target:
                replacement = ColumnInfo(target, FIELD_DEFINITIONS[target]["label"], column.original or f"Inferred {target}", 0.86, "inferred")
        if replacement:
            for record in records:
                record[replacement.key] = record.pop(column.key, "")
            occupied.add(replacement.key)
            updated.append(replacement)
        else:
            updated.append(column)
    return updated


def clean_student_records(
    records: list[dict[str, Any]], columns: list[ColumnInfo], dataset_type: str
) -> tuple[list[dict[str, Any]], dict[str, int]]:
    diagnostics = {"blank_rows_removed": 0, "non_student_rows_removed": 0, "gender_suffix_split": 0}
    name_key = find_base_key(columns, "student_name")
    row_key = find_base_key(columns, "row_number")
    gender_key = find_base_key(columns, "gender")
    if dataset_type not in {"student_roster", "promotion_table", "monitoring_table", "assessment_results"} or not name_key:
        return records, diagnostics

    cleaned: list[dict[str, Any]] = []
    for record in records:
        name = str(record.get(name_key, "") or "").strip()
        if not name:
            diagnostics["blank_rows_removed"] += 1
            continue
        if is_footer_or_label(name) or recognize_header(name).key == "student_name":
            diagnostics["non_student_rows_removed"] += 1
            continue
        # Lists and monitoring files should contain a plausible person name.
        if dataset_type in {"student_roster", "promotion_table"} and not looks_like_person_name(name):
            diagnostics["non_student_rows_removed"] += 1
            continue

        suffix_match = re.match(r"^(.*?)(?:\s+)(erkak|ayol|male|female|мужской|женский)$", name, re.I)
        if suffix_match:
            record[name_key] = suffix_match.group(1).strip()
            if not gender_key:
                gender_key = "gender"
                columns.append(ColumnInfo("gender", "Gender", "Inferred from name suffix", 0.85, "inferred"))
            if not cell_has_value(record.get(gender_key, "")):
                record[gender_key] = suffix_match.group(2)
            diagnostics["gender_suffix_split"] += 1

        if row_key and cell_has_value(record.get(row_key, "")) and not looks_like_sequence(record.get(row_key)):
            # Keep unusual numbering, but reject another textual header masquerading as data.
            if recognize_header(record.get(row_key)).method != "custom":
                diagnostics["non_student_rows_removed"] += 1
                continue
        cleaned.append(record)
    return cleaned, diagnostics


def records_from_explicit_headers(
    headers: list[Any],
    data_rows: list[list[Any]],
    source_file: str,
    source_sheet: str,
    start_row: int = 1,
) -> tuple[list[ColumnInfo], list[dict[str, Any]]]:
    combined, _ = collapse_duplicate_columns([headers, *data_rows])
    if not combined:
        return [], []
    headers = combined[0]
    data_rows = combined[1:]
    columns = make_unique_columns(headers)
    records: list[dict[str, Any]] = []
    for offset, row in enumerate(data_rows, start=start_row):
        if count_non_empty(row) == 0 or is_repeated_header_row(row, headers):
            continue
        record: dict[str, Any] = {
            "__source_file": source_file,
            "__source_sheet": source_sheet,
            "__source_row": offset,
        }
        meaningful = 0
        for index, column in enumerate(columns):
            value = normalize_cell_value(row[index] if index < len(row) else "")
            record[column.key] = value
            meaningful += int(cell_has_value(value))
        if meaningful:
            records.append(record)
        if len(records) >= MAX_ROWS_PER_FILE:
            break
    return columns, records


def classify_dataset(columns: list[ColumnInfo], records: list[dict[str, Any]], context_text: str = "") -> tuple[str, float, list[str]]:
    keys = {canonical_base_key(column.key) for column in columns}
    normalized_context = normalize_text(context_text)
    reasons: list[str] = []

    if "promoted_class" in keys:
        reasons.append("current and promoted class fields")
        return "promotion_table", 0.99, reasons
    if any(word in normalized_context for word in ("saldo", "сальдо", "debet", "дебет", "kredit", "кредит", "akt sverka", "акт свер")):
        reasons.append("financial reconciliation vocabulary")
        return "financial_table", 0.98, reasons
    if "student_name" in keys:
        question_count = sum(1 for key in keys if key.startswith("question_"))
        if (
            "monitoring" in normalized_context
            or question_count >= 2
            or (question_count >= 1 and ("total" in keys or "percentage" in keys))
            or {"total", "percentage"} <= keys
        ):
            reasons.append("student names with monitoring question/result columns")
            return "monitoring_table", 0.97, reasons
        if "subject" in keys or "place" in keys or "score" in keys:
            reasons.append("student names with assessment result fields")
            return "assessment_results", 0.94, reasons
        reasons.append("student-name column")
        return "student_roster", 0.96, reasons
    if "monitoring" in normalized_context and len(columns) >= 3:
        reasons.append("monitoring title")
        return "monitoring_table", 0.82, reasons
    if len(columns) >= 2 and records:
        reasons.append("structured table")
        return "generic_table", 0.65, reasons
    return "unknown_table", 0.35, ["insufficient structure"]


def table_signature_v06(columns: list[ColumnInfo], dataset_type: str) -> tuple[str, ...]:
    keys: list[str] = []
    for column in columns:
        base = canonical_base_key(column.key)
        if base.startswith("custom_"):
            base = f"custom:{normalize_text(column.original)}"
        keys.append(base)
    # Rosters intentionally use a union schema so simple and rich class lists can be combined.
    if dataset_type == "student_roster":
        return ("student_roster",)
    if dataset_type == "promotion_table":
        return ("promotion_table",)
    return (dataset_type, *sorted(keys))


def one_column_roster(rows: list[list[Any]], source_file: str, source_sheet: str, context_text: str) -> dict[str, Any] | None:
    if not rows or max((len(row) for row in rows), default=0) != 1:
        return None
    values = [str(row[0] or "").strip() for row in rows if row and cell_has_value(row[0])]
    if len(values) < 3:
        return None
    context_normalized = normalize_text(f"{context_text[:600]} {source_file} {source_sheet}")
    class_hint = bool(extract_class_values(context_normalized))
    roster_hint = any(token in context_normalized for token in ("ro'yxat", "ro yxat", "ro`yxat", "sinf o'quvchi", "sinf oquvchi", "class list"))
    blocked_context = any(token in context_normalized for token in ("metodik tavsiya", "loyiha ishi", "annotatsiya", "mundarija", "bayonnoma"))
    explicit_roster_hint = any(token in context_normalized for token in ("ro'yxat", "ro yxat", "ro`yxat", "class list"))
    if blocked_context and not explicit_roster_hint:
        return None
    header_index = next((i for i, value in enumerate(values[:10]) if recognize_header(value).key == "student_name"), None)
    start = (header_index + 1) if header_index is not None else 0
    names = [value for value in values[start:] if looks_like_person_name(value) and not is_footer_or_label(value)]
    if len(names) < 2:
        return None
    columns = [
        ColumnInfo("row_number", "No.", "Generated sequence", 0.98, "inferred"),
        ColumnInfo("student_name", "Student name", values[header_index] if header_index is not None else "Inferred student name", 0.95, "inferred"),
    ]
    records = [
        {
            "__source_file": source_file,
            "__source_sheet": source_sheet,
            "__source_row": start + index + 1,
            "row_number": index,
            "student_name": name,
        }
        for index, name in enumerate(names, start=1)
    ]
    class_info = infer_class([(context_text, "document title", 100), (source_file, "file name", 80), (source_sheet, "sheet name", 40)])
    columns, records = apply_inferred_class(columns, records, class_info)
    return {
        "sheet": source_sheet,
        "header_index": header_index if header_index is not None else -1,
        "header_mode": "single_column_roster",
        "columns": columns,
        "records": records,
        "class_info": class_info,
        "diagnostics": {"generated_row_numbers": len(records)},
    }


def headerless_semantic_table(rows: list[list[Any]], source_file: str, source_sheet: str, context_text: str) -> dict[str, Any] | None:
    rows, duplicate_columns = collapse_duplicate_columns(rows)
    if len(rows) < 2:
        return None
    width = max((len(row) for row in rows), default=0)
    sample = rows[: min(25, len(rows))]
    if width < 2:
        return None

    def ratio(column: int, predicate: Any) -> float:
        relevant = [row[column] if column < len(row) else "" for row in sample]
        return sum(predicate(value) for value in relevant) / max(1, len(relevant))

    if ratio(0, looks_like_sequence) >= 0.6:
        if width >= 4 and ratio(1, looks_like_class) >= 0.55 and ratio(2, looks_like_person_name) >= 0.55 and ratio(3, looks_like_class) >= 0.55:
            headers = ["№", "Sinf", "O‘quvchi", "Ko'chirilgan sinfi"] + [f"Column {i + 1}" for i in range(4, width)]
            mode = "inferred_promotion_table"
        else:
            person_column = max(range(1, width), key=lambda index: ratio(index, looks_like_person_name))
            if ratio(person_column, looks_like_person_name) < 0.55:
                return None
            headers = [f"Column {i + 1}" for i in range(width)]
            headers[0] = "№"
            headers[person_column] = "O‘quvchi F.I.Sh."
            for index in range(width):
                if index not in {0, person_column} and ratio(index, looks_like_class) >= 0.55:
                    headers[index] = "Sinf" if "Sinf" not in headers else "Ko'chirilgan sinfi"
            mode = "inferred_roster"
        columns, records = records_from_explicit_headers(headers, rows, source_file, source_sheet, start_row=1)
        class_info = infer_class([(context_text, "document title", 100), (source_file, "file name", 80), (source_sheet, "sheet name", 40)])
        columns, records = apply_inferred_class(columns, records, class_info)
        return {
            "sheet": source_sheet,
            "header_index": -1,
            "header_mode": mode,
            "columns": columns,
            "records": records,
            "class_info": class_info,
            "diagnostics": {"duplicate_columns_removed": duplicate_columns},
        }
    return None


def strong_explicit_header_index(rows: list[list[Any]]) -> int | None:
    sparse = explicit_roster_header_index(rows)
    if sparse is not None:
        return sparse
    for index, row in enumerate(rows[:20]):
        infos = [recognize_header(value) for value in row]
        # Numeric headings are valid in a monitoring header only when anchored by
        # semantic labels such as No. and Student name. A normal data row full of
        # scores must never become a header merely because numbers map to
        # question_N fields.
        semantic = [info for info in infos if info.method not in {"custom", "empty", "question"}]
        keys = {info.key for info in semantic}
        question_count = sum(info.method == "question" for info in infos)
        if len(keys) >= 2:
            return index
        if "student_name" in keys and ("row_number" in keys or question_count >= 2):
            return index
        if question_count >= 3 and {"row_number", "student_name"}.issubset(keys):
            return index
    return None


def dataset_from_rows(
    rows: list[list[Any]],
    source_file: str,
    source_sheet: str,
    context_text: str = "",
    source_kind: str = "spreadsheet",
) -> dict[str, Any] | None:
    rows, duplicate_columns = collapse_duplicate_columns(rows)
    rows = [row for row in rows if count_non_empty(row)]
    if not rows:
        return None

    special = one_column_roster(rows, source_file, source_sheet, context_text)
    explicit_index = None if special is not None else strong_explicit_header_index(rows)
    if special is None and explicit_index is None:
        special = headerless_semantic_table(rows, source_file, source_sheet, context_text)

    if special is not None:
        columns = special["columns"]
        records = special["records"]
        header_index = special["header_index"]
        header_mode = special["header_mode"]
        class_info = special["class_info"]
        diagnostics = dict(special.get("diagnostics", {}))
    else:
        header_index = explicit_index if explicit_index is not None else detect_header_row(rows)
        columns, records = records_from_explicit_headers(
            rows[header_index], rows[header_index + 1 :], source_file, source_sheet, start_row=header_index + 2
        )
        preheader = " ".join(str(value) for row in rows[:header_index] for value in row if cell_has_value(value))
        class_info = infer_class([
            (preheader or context_text, "document title", 100),
            (source_file, "file name", 80),
            (source_sheet, "sheet name", 40),
        ])
        header_mode = "explicit"
        diagnostics = {}

    # Recover blank or unfamiliar headings from their values before injecting a
    # class inferred from the title/filename. Otherwise a genuine class column
    # headed "Sinflar" could be mistaken for promoted_class after the inferred
    # constant class field had already occupied the canonical class key.
    columns = infer_columns_from_values(columns, records)
    columns, records = apply_inferred_class(columns, records, class_info)
    records, merged_count = merge_continuation_records(records, columns)
    dataset_type, confidence, reasons = classify_dataset(columns, records, f"{context_text} {source_sheet} {source_file}")
    records, cleanup = clean_student_records(records, columns, dataset_type)
    diagnostics.update(cleanup)
    diagnostics["duplicate_columns_removed"] = diagnostics.get("duplicate_columns_removed", 0) + duplicate_columns
    diagnostics["continuation_rows_merged"] = merged_count

    if records:
        visible_keys = [column.key for column in columns]
        average_density = sum(
            sum(cell_has_value(record.get(key, "")) for key in visible_keys)
            for record in records
        ) / len(records)
        diagnostics["average_nonempty_cells"] = round(average_density, 2)
        has_student_name = find_base_key(columns, "student_name") is not None
        if average_density <= 1.2 and dataset_type in {"monitoring_table", "generic_table", "unknown_table"}:
            return None
        if dataset_type == "monitoring_table" and not has_student_name and average_density <= 2.2:
            return None

    if not records:
        return None
    return {
        "name": source_sheet,
        "sheet": source_sheet,
        "header_index": header_index,
        "header_mode": header_mode,
        "columns": columns,
        "records": records,
        "class_info": class_info,
        "dataset_type": dataset_type,
        "dataset_confidence": confidence,
        "classification_reasons": reasons,
        "signature": table_signature_v06(columns, dataset_type),
        "source_file": source_file,
        "source_kind": source_kind,
        "diagnostics": diagnostics,
    }


def classify_document(filename: str, text: str, datasets: list[dict[str, Any]]) -> dict[str, Any]:
    normalized = normalize_text(f"{filename} {text}")
    scores: defaultdict[str, float] = defaultdict(float)
    reasons: defaultdict[str, list[str]] = defaultdict(list)

    def add(kind: str, score: float, reason: str) -> None:
        scores[kind] += score
        reasons[kind].append(reason)

    keyword_rules = {
        "financial_document": [("akt sverka", 6), ("акт свер", 6), ("saldo", 4), ("сальдо", 4), ("debet", 3), ("дебет", 3), ("kredit", 3), ("кредит", 3)],
        "meeting_minutes": [("bayonnoma", 5), ("yig'ilish", 4), ("yig ilish", 4), ("kun tartibi", 4), ("ped kengash", 5), ("qaror qiladi", 3)],
        "methodical_document": [("metodik tavsiya", 6), ("loyiha ishi", 5), ("annotatsiya", 3), ("mundarija", 3), ("asosiy qism", 1)],
        "assessment_material": [("test", 3), ("dars", 1), ("savol", 2), ("qaysi", 1)],
        "monitoring_document": [("monitoring", 7), ("samaradorlik", 2)],
        "student_roster_document": [("ro'yxat", 5), ("ro yxat", 5), ("ro`yxat", 5), ("o'quvchilari", 2), ("o‘quvchilari", 2)],
        "promotion_document": [("ko'chirilgan sinfi", 7), ("ko chirilgan sinfi", 7), ("sinfdan sinfga", 3)],
    }
    for kind, rules in keyword_rules.items():
        for keyword, score in rules:
            if keyword in normalized:
                add(kind, score, f"keyword: {keyword}")

    # Narrative container types should remain the document classification even
    # when they embed a large structured appendix. For example, pedagogical
    # meeting minutes can contain hundreds of promotion rows but are still
    # meeting minutes as a whole.
    if scores["meeting_minutes"] >= 10:
        add("meeting_minutes", 4, "strong meeting-minutes structure")
    if scores["methodical_document"] >= 8:
        add("methodical_document", 3, "strong methodical-document structure")
    if scores["financial_document"] >= 8:
        add("financial_document", 3, "strong financial-document structure")

    question_count = len(re.findall(r"(?:^|\n)\s*\d{1,3}[.)]\s+", text))
    if question_count >= 5:
        add("assessment_material", min(8, question_count / 2), f"{question_count} numbered questions")

    # Aggregate table evidence by type. A document containing twenty-nine
    # promotion tables and one ordinary roster should be classified as a
    # promotion document, not decided by whichever table type appeared first.
    mapping = {
        "student_roster": "student_roster_document",
        "promotion_table": "promotion_document",
        "monitoring_table": "monitoring_document",
        "assessment_results": "assessment_results_document",
        "financial_table": "financial_document",
        "generic_table": "data_document",
    }
    dataset_stats: defaultdict[str, dict[str, float]] = defaultdict(lambda: {"count": 0.0, "rows": 0.0})
    for dataset in datasets:
        kind = mapping.get(dataset.get("dataset_type"), "data_document")
        dataset_stats[kind]["count"] += 1
        dataset_stats[kind]["rows"] += len(dataset.get("records", []))

    total_tables = max(1.0, float(len(datasets)))
    total_rows = max(1.0, sum(stats["rows"] for stats in dataset_stats.values()))
    for kind, stats in dataset_stats.items():
        table_share = stats["count"] / total_tables
        row_share = stats["rows"] / total_rows
        evidence = 5.5 + min(4.0, stats["rows"] / 80.0) + table_share * 5.0 + row_share * 2.0
        add(
            kind,
            evidence,
            f"{int(stats['count'])} detected table(s), {int(stats['rows'])} row(s)",
        )

    if not scores:
        if datasets:
            return {"type": "data_document", "confidence": 0.65, "reasons": ["contains structured tables"]}
        return {"type": "unknown_document", "confidence": 0.35, "reasons": ["no strong document pattern"]}

    ranked = sorted(scores.items(), key=lambda item: item[1], reverse=True)
    winner, winning_score = ranked[0]
    runner_up = ranked[1][1] if len(ranked) > 1 else 0
    confidence = min(0.99, 0.58 + winning_score * 0.035 + max(0, winning_score - runner_up) * 0.02)
    return {"type": winner, "confidence": round(confidence, 2), "reasons": reasons[winner][:5]}


def finalize_parsed_file(
    filename: str,
    detected_type: str,
    source_kind: str,
    text: str,
    datasets: list[dict[str, Any]],
    warnings: list[str] | None = None,
    extra: dict[str, Any] | None = None,
) -> dict[str, Any]:
    classification = classify_document(filename, text, datasets)
    result: dict[str, Any] = {
        "detected_type": detected_type,
        "source_kind": source_kind,
        "document_type": classification["type"],
        "document_confidence": classification["confidence"],
        "classification_reasons": classification["reasons"],
        "datasets": datasets,
        "warnings": warnings or [],
        "document_text_excerpt": re.sub(r"\s+", " ", text).strip()[:500],
    }
    if extra:
        result.update(extra)
    if datasets:
        primary = max(datasets, key=lambda item: (item.get("dataset_confidence", 0), len(item.get("records", []))))
        result.update({
            "sheet": primary["sheet"],
            "header_index": primary["header_index"],
            "header_mode": primary.get("header_mode", "automatic"),
            "header_rows": primary.get("header_rows", 1),
            "columns": primary["columns"],
            "records": primary["records"],
            "class_info": primary.get("class_info", {}),
            "continuation_rows_merged": primary.get("diagnostics", {}).get("continuation_rows_merged", 0),
        })
    else:
        result.update({"sheet": "—", "header_index": -1, "header_mode": "none", "columns": [], "records": [], "class_info": {}})
    return result


def parse_xlsx(content: bytes, filename: str) -> dict[str, Any]:
    workbook = load_workbook(io.BytesIO(content), data_only=True, read_only=False)
    datasets: list[dict[str, Any]] = []
    warnings: list[str] = []
    title_fragments: list[str] = []
    for sheet in workbook.worksheets:
        rows = [list(row) for row in sheet.iter_rows(values_only=True)]
        if not any(count_non_empty(row) for row in rows):
            continue
        title = " ".join(str(value) for row in rows[:5] for value in row if cell_has_value(value))
        title_fragments.append(title)
        try:
            dataset = dataset_from_rows(rows, filename, sheet.title, title, "spreadsheet")
            if dataset:
                datasets.append(dataset)
            else:
                warnings.append(f"Sheet '{sheet.title}' did not contain a reliable data table.")
        except Exception as exc:
            warnings.append(f"Sheet '{sheet.title}': {exc}")
    if not datasets and not warnings:
        raise ValueError("Workbook contains no usable worksheet.")
    return finalize_parsed_file(filename, "XLSX workbook", "spreadsheet", " ".join(title_fragments), datasets, warnings, {"sheet_count": len(workbook.worksheets)})


def parse_xls(content: bytes, filename: str) -> dict[str, Any]:
    if xlrd is None:
        raise ValueError("Old .xls support requires xlrd. Run: pip install -r requirements.txt")
    workbook = xlrd.open_workbook(file_contents=content)
    datasets: list[dict[str, Any]] = []
    warnings: list[str] = []
    texts: list[str] = []
    for index in range(workbook.nsheets):
        sheet = workbook.sheet_by_index(index)
        if not sheet.nrows:
            continue
        rows = [sheet.row_values(i) for i in range(sheet.nrows)]
        title = " ".join(str(value) for row in rows[:5] for value in row if cell_has_value(value))
        texts.append(title)
        try:
            dataset = dataset_from_rows(rows, filename, sheet.name, title, "spreadsheet")
            if dataset:
                datasets.append(dataset)
            else:
                warnings.append(f"Sheet '{sheet.name}' did not contain a reliable data table.")
        except Exception as exc:
            warnings.append(f"Sheet '{sheet.name}': {exc}")
    return finalize_parsed_file(filename, "Legacy binary XLS workbook", "spreadsheet", " ".join(texts), datasets, warnings, {"sheet_count": workbook.nsheets})


def parse_csv(content: bytes, filename: str) -> dict[str, Any]:
    text = decode_csv(content)
    sample = text[:4096]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
        delimiter = dialect.delimiter
    except csv.Error:
        delimiter = ","
    rows = [row for row in csv.reader(io.StringIO(text), delimiter=delimiter) if any(cell.strip() for cell in row)]
    dataset = dataset_from_rows(rows, filename, "CSV", filename, "spreadsheet")
    return finalize_parsed_file(filename, "CSV text", "spreadsheet", filename, [dataset] if dataset else [], [])


def html_table_to_matrix(table: Any) -> tuple[list[list[str]], list[str]]:
    source_rows = table.find_all("tr")
    grid: dict[tuple[int, int], str] = {}
    max_column = 0
    row_types: list[str] = []
    for row_index, tr in enumerate(source_rows):
        cells = tr.find_all(["th", "td"], recursive=False)
        row_types.append("header" if cells and all(cell.name == "th" for cell in cells) else "data")
        column_index = 0
        for cell in cells:
            while (row_index, column_index) in grid:
                column_index += 1
            value = unescape(cell.get_text(" ", strip=True))
            rowspan = max(1, int(cell.get("rowspan", 1) or 1))
            colspan = max(1, int(cell.get("colspan", 1) or 1))
            for target_row in range(row_index, row_index + rowspan):
                for target_column in range(column_index, column_index + colspan):
                    grid[(target_row, target_column)] = value
                    max_column = max(max_column, target_column + 1)
            column_index += colspan
    matrix = [[grid.get((r, c), "") for c in range(max_column)] for r in range(len(source_rows))]
    return matrix, row_types


def dataset_from_html_matrix(
    matrix: list[list[Any]],
    row_types: list[str],
    filename: str,
    source_sheet: str,
    context_text: str,
) -> dict[str, Any] | None:
    matrix, duplicate_columns = collapse_duplicate_columns(matrix)
    if not matrix:
        return None
    first_header = next((index for index, kind in enumerate(row_types) if kind == "header"), None)
    if first_header is None:
        dataset = dataset_from_rows(matrix, filename, source_sheet, context_text, "spreadsheet")
        if dataset:
            dataset["diagnostics"]["duplicate_columns_removed"] = dataset["diagnostics"].get("duplicate_columns_removed", 0) + duplicate_columns
        return dataset

    last_header = first_header
    while last_header + 1 < len(row_types) and row_types[last_header + 1] == "header":
        last_header += 1
    headers: list[str] = []
    width = max((len(row) for row in matrix), default=0)
    for column_index in range(width):
        parts: list[str] = []
        for row_index in range(first_header, last_header + 1):
            value = str(matrix[row_index][column_index] if column_index < len(matrix[row_index]) else "").strip()
            if value and (not parts or normalize_text(value) != normalize_text(parts[-1])):
                parts.append(value)
        headers.append(" / ".join(parts) if parts else f"Column {column_index + 1}")

    columns, records = records_from_explicit_headers(
        headers,
        matrix[last_header + 1 :],
        filename,
        source_sheet,
        start_row=last_header + 2,
    )
    class_info = infer_class([(context_text, "document title", 100), (filename, "file name", 80), (source_sheet, "sheet name", 40)])
    columns, records = apply_inferred_class(columns, records, class_info)
    columns = infer_columns_from_values(columns, records)
    records, merged_count = merge_continuation_records(records, columns)
    dataset_type, confidence, reasons = classify_dataset(columns, records, f"{context_text} {filename}")
    records, cleanup = clean_student_records(records, columns, dataset_type)
    if not records:
        return None
    diagnostics = dict(cleanup)
    diagnostics.update({
        "duplicate_columns_removed": duplicate_columns,
        "continuation_rows_merged": merged_count,
        "header_rows_combined": last_header - first_header + 1,
    })
    return {
        "name": source_sheet,
        "sheet": source_sheet,
        "header_index": last_header,
        "header_mode": "multi_row" if last_header > first_header else "explicit",
        "header_rows": last_header - first_header + 1,
        "columns": columns,
        "records": records,
        "class_info": class_info,
        "dataset_type": dataset_type,
        "dataset_confidence": confidence,
        "classification_reasons": reasons,
        "signature": table_signature_v06(columns, dataset_type),
        "source_file": filename,
        "source_kind": "spreadsheet",
        "diagnostics": diagnostics,
    }


def parse_html_table(content: bytes, filename: str) -> dict[str, Any]:
    text = decode_csv(content)
    soup = BeautifulSoup(text, "html.parser")
    tables = soup.find_all("table")
    if not tables:
        linked = [link.get("href") for link in soup.find_all("link") if link.get("href")]
        if linked:
            warning = "This Excel HTML file is only an index and references missing sidecar files: " + ", ".join(linked[:3])
            return finalize_parsed_file(filename, "Excel HTML frameset", "spreadsheet", soup.get_text(" ", strip=True), [], [warning])
        raise ValueError("The file looks like HTML, but no table was found.")
    datasets: list[dict[str, Any]] = []
    warnings: list[str] = []
    context = soup.title.get_text(" ", strip=True) if soup.title else filename
    for index, table in enumerate(tables, start=1):
        matrix, row_types = html_table_to_matrix(table)
        if not matrix:
            continue
        try:
            dataset = dataset_from_html_matrix(matrix, row_types, filename, f"HTML table {index}", context)
            if dataset:
                datasets.append(dataset)
        except Exception as exc:
            warnings.append(f"HTML table {index}: {exc}")
    return finalize_parsed_file(filename, "HTML table saved as .xls", "spreadsheet", soup.get_text(" ", strip=True)[:5000], datasets, warnings, {"table_count": len(tables)})


def parse_spreadsheetml(content: bytes, filename: str) -> dict[str, Any]:
    try:
        root = ET.fromstring(content)
    except ET.ParseError as exc:
        raise ValueError(f"Invalid SpreadsheetML XML: {exc}") from exc
    ns = {"ss": "urn:schemas-microsoft-com:office:spreadsheet"}
    worksheets = root.findall(".//ss:Worksheet", ns) or [node for node in root.iter() if node.tag.endswith("Worksheet")]
    datasets: list[dict[str, Any]] = []
    warnings: list[str] = []
    for worksheet in worksheets:
        sheet_name = worksheet.attrib.get("{urn:schemas-microsoft-com:office:spreadsheet}Name") or worksheet.attrib.get("Name") or "SpreadsheetML"
        row_nodes = worksheet.findall(".//ss:Row", ns) or [node for node in worksheet.iter() if node.tag.endswith("Row")]
        rows: list[list[Any]] = []
        for row_node in row_nodes:
            row: list[Any] = []
            cell_nodes = list(row_node.findall("ss:Cell", ns)) or [node for node in row_node if node.tag.endswith("Cell")]
            for cell in cell_nodes:
                index_value = cell.attrib.get("{urn:schemas-microsoft-com:office:spreadsheet}Index") or cell.attrib.get("Index")
                if index_value:
                    while len(row) < int(index_value) - 1:
                        row.append("")
                data_node = cell.find("ss:Data", ns)
                if data_node is None:
                    data_node = next((node for node in cell.iter() if node.tag.endswith("Data")), None)
                row.append(data_node.text if data_node is not None and data_node.text is not None else "")
            if row:
                rows.append(row)
        try:
            dataset = dataset_from_rows(rows, filename, sheet_name, sheet_name, "spreadsheet")
            if dataset:
                datasets.append(dataset)
        except Exception as exc:
            warnings.append(f"Sheet '{sheet_name}': {exc}")
    return finalize_parsed_file(filename, "SpreadsheetML XML saved as .xls", "spreadsheet", filename, datasets, warnings, {"sheet_count": len(worksheets)})


def word_table_to_rows(table: Any) -> list[list[Any]]:
    rows: list[list[Any]] = []
    for table_row in table.rows:
        values = [re.sub(r"\s+", " ", cell.text or "").strip() for cell in table_row.cells]
        if any(cell_has_value(value) for value in values):
            rows.append(values)
    cleaned, _ = collapse_duplicate_columns(rows)
    return cleaned


def docx_xml_fallback(content: bytes) -> tuple[str, list[list[list[str]]]]:
    """Read paragraphs and tables directly from document.xml, ignoring broken media."""
    with zipfile.ZipFile(io.BytesIO(content)) as archive:
        xml = archive.read("word/document.xml")
    root = ET.fromstring(xml)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    def node_text(node: Any) -> str:
        values = [text.text or "" for text in node.findall(".//w:t", ns)]
        return re.sub(r"\s+", " ", "".join(values)).strip()

    text = " ".join(node_text(paragraph) for paragraph in root.findall(".//w:body/w:p", ns) if node_text(paragraph))
    tables: list[list[list[str]]] = []
    for table in root.findall(".//w:tbl", ns):
        rows: list[list[str]] = []
        for tr in table.findall("./w:tr", ns):
            row: list[str] = []
            for tc in tr.findall("./w:tc", ns):
                value = node_text(tc)
                grid_span = tc.find("./w:tcPr/w:gridSpan", ns)
                span = int(grid_span.attrib.get(f"{{{ns['w']}}}val", "1")) if grid_span is not None else 1
                row.append(value)
                row.extend([""] * max(0, span - 1))
            if any(cell_has_value(value) for value in row):
                rows.append(row)
        rows, _ = collapse_duplicate_columns(rows)
        if rows:
            tables.append(rows)
    return text, tables


def parse_docx(content: bytes, filename: str) -> dict[str, Any]:
    warnings: list[str] = []
    try:
        document = Document(io.BytesIO(content))
        document_text = " ".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
        raw_tables = [word_table_to_rows(table) for table in document.tables]
        detected_type = "DOCX Word document"
    except Exception as exc:
        try:
            document_text, raw_tables = docx_xml_fallback(content)
            warnings.append(f"Word media was damaged; table text was recovered from document.xml ({exc}).")
            detected_type = "DOCX recovered from XML"
        except Exception as fallback_exc:
            raise ValueError(f"Could not open the DOCX document: {exc}. XML recovery also failed: {fallback_exc}") from fallback_exc

    datasets: list[dict[str, Any]] = []
    table_summaries: list[dict[str, Any]] = []
    for index, rows in enumerate(raw_tables, start=1):
        summary: dict[str, Any] = {"table": index, "raw_rows": len(rows), "raw_columns": max((len(row) for row in rows), default=0)}
        try:
            dataset = dataset_from_rows(rows, filename, f"Word table {index}", document_text, "word")
            if dataset:
                dataset["table_index"] = index
                datasets.append(dataset)
                summary.update({
                    "usable": True,
                    "rows": len(dataset["records"]),
                    "columns": len(dataset["columns"]),
                    "dataset_type": dataset["dataset_type"],
                    "headers": [column.original for column in dataset["columns"]],
                    "header_mode": dataset["header_mode"],
                    "diagnostics": dataset.get("diagnostics", {}),
                })
            else:
                summary.update({"usable": False, "error": "No reliable data rows"})
        except Exception as exc:
            summary.update({"usable": False, "error": str(exc)})
        table_summaries.append(summary)

    return finalize_parsed_file(
        filename,
        detected_type,
        "word",
        document_text,
        datasets,
        warnings,
        {"table_count": len(raw_tables), "usable_table_count": len(datasets), "table_summaries": table_summaries},
    )


def parse_pdf(content: bytes, filename: str) -> dict[str, Any]:
    if pdfplumber is None:
        raise ValueError("PDF support requires pdfplumber. Run: pip install -r requirements.txt")
    datasets: list[dict[str, Any]] = []
    warnings: list[str] = []
    text_parts: list[str] = []
    table_count = 0
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page_index, page in enumerate(pdf.pages[:100], start=1):
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
            try:
                tables = page.extract_tables() or []
            except Exception as exc:
                warnings.append(f"Page {page_index}: table extraction failed ({exc}).")
                tables = []
            for local_index, rows in enumerate(tables, start=1):
                table_count += 1
                cleaned = [[cell or "" for cell in row] for row in rows if row]
                try:
                    dataset = dataset_from_rows(cleaned, filename, f"PDF page {page_index} table {local_index}", page_text, "pdf")
                    if dataset:
                        dataset["page"] = page_index
                        datasets.append(dataset)
                except Exception as exc:
                    warnings.append(f"Page {page_index} table {local_index}: {exc}")
    text = "\n".join(text_parts)
    if not text.strip() and not datasets:
        warnings.append("No selectable text or reliable tables were found. This PDF may be scanned.")
    return finalize_parsed_file(filename, "PDF document", "pdf", text, datasets, warnings, {"table_count": table_count})


def parse_uploaded_file(filename: str, content: bytes) -> dict[str, Any]:
    if not content:
        raise ValueError("The selected file is empty (0 bytes).")
    suffix = Path(filename).suffix.lower()
    detected = sniff_file_type(content, filename)

    if suffix == ".pdf" or content.startswith(b"%PDF"):
        return parse_pdf(content, filename)
    if detected == "docx-zip":
        return parse_docx(content, filename)
    if detected == "xlsx-zip":
        result = parse_xlsx(content, filename)
        if suffix not in {".xlsx", ".xlsm"}:
            result["detected_type"] += " (renamed extension)"
        return result
    if detected == "ole2":
        if suffix == ".doc":
            return parse_doc(content, filename)
        return parse_xls(content, filename)
    if detected == "html-table":
        return parse_html_table(content, filename)
    if detected == "spreadsheetml-xml":
        return parse_spreadsheetml(content, filename)
    if detected == "delimited-text":
        return parse_csv(content, filename)

    try:
        if suffix == ".docx":
            return parse_docx(content, filename)
        if suffix == ".doc":
            return parse_doc(content, filename)
        if suffix in {".xlsx", ".xlsm"}:
            return parse_xlsx(content, filename)
        if suffix == ".xls":
            return parse_xls(content, filename)
        if suffix == ".csv":
            return parse_csv(content, filename)
    except Exception as exc:
        prefix = content[:12].hex(" ").upper()
        raise ValueError(
            f"Could not read this file as {suffix or 'a supported document'}. "
            f"Detected content: {detected}. First bytes: {prefix}. Original error: {exc}"
        ) from exc

    raise ValueError(
        f"Unsupported file. Extension: {suffix or 'none'}; detected content: {detected}. "
        "Supported inputs are XLSX, XLS, XLSM, CSV, DOCX, DOC, and PDF."
    )


def merge_parsed_files(parsed_files: list[dict[str, Any]]) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    column_map: OrderedDict[str, dict[str, Any]] = OrderedDict()
    records: list[dict[str, Any]] = []
    core_default = {"row_number", "student_name", "class", "grade", "promoted_class", "total", "percentage"}
    sensitive = {"pinfl", "birth_date", "phone", "email", "address", "parent_phone", "parent_email"}
    for parsed in parsed_files:
        for column in parsed["columns"]:
            base_key = canonical_base_key(column.key)
            if column.key not in column_map:
                column_map[column.key] = {
                    "key": column.key,
                    "label": column.label,
                    "originals": [column.original],
                    "confidence": column.confidence,
                    "method": column.method,
                    "default_selected": base_key in core_default or column.method == "question",
                    "sensitive": base_key in sensitive,
                }
            else:
                existing = column_map[column.key]
                if column.original not in existing["originals"]:
                    existing["originals"].append(column.original)
                existing["confidence"] = max(existing["confidence"], column.confidence)
        records.extend(parsed["records"])
    columns = list(column_map.values())
    if columns and not any(column["default_selected"] for column in columns):
        for column in columns[: min(4, len(columns))]:
            column["default_selected"] = True
    return columns, records


def group_datasets(datasets: list[dict[str, Any]]) -> list[dict[str, Any]]:
    buckets: OrderedDict[tuple[str, ...], list[dict[str, Any]]] = OrderedDict()
    for dataset in datasets:
        signature = tuple(dataset.get("signature") or table_signature_v06(dataset["columns"], dataset["dataset_type"]))
        buckets.setdefault(signature, []).append(dataset)

    groups: list[dict[str, Any]] = []
    labels = {
        "student_roster": "Student rosters",
        "promotion_table": "Class promotion tables",
        "monitoring_table": "Monitoring tables",
        "assessment_results": "Assessment results",
        "financial_table": "Financial tables",
        "generic_table": "Other data tables",
        "unknown_table": "Unclassified tables",
    }
    priorities = {
        "promotion_table": 100,
        "student_roster": 95,
        "monitoring_table": 80,
        "assessment_results": 70,
        "financial_table": 40,
        "generic_table": 30,
        "unknown_table": 10,
    }
    for index, (signature, items) in enumerate(buckets.items(), start=1):
        columns, records = merge_parsed_files(items)
        dataset_type = items[0]["dataset_type"]
        sources = sorted({item["source_file"] for item in items})
        group_id = hashlib.sha1(("|".join(signature) + "|" + "|".join(sources)).encode("utf-8")).hexdigest()[:12]
        groups.append({
            "id": group_id,
            "name": labels.get(dataset_type, dataset_type.replace("_", " ").title()),
            "dataset_type": dataset_type,
            "columns": columns,
            "records": records,
            "row_count": len(records),
            "source_count": len(sources),
            "sources": sources,
            "table_count": len(items),
            "signature": list(signature),
            "priority": priorities.get(dataset_type, 0),
            "diagnostics": {
                "duplicate_columns_removed": sum(item.get("diagnostics", {}).get("duplicate_columns_removed", 0) for item in items),
                "blank_rows_removed": sum(item.get("diagnostics", {}).get("blank_rows_removed", 0) for item in items),
                "non_student_rows_removed": sum(item.get("diagnostics", {}).get("non_student_rows_removed", 0) for item in items),
                "continuation_rows_merged": sum(item.get("diagnostics", {}).get("continuation_rows_merged", 0) for item in items),
            },
        })
    groups.sort(key=lambda group: (group["priority"], group["row_count"], group["source_count"]), reverse=True)
    return groups


class SelectDatasetRequest(BaseModel):
    session_id: str
    dataset_id: str


@app.get("/", response_class=HTMLResponse)
def index() -> FileResponse:
    return FileResponse(STATIC_DIR / "index.html")


@app.get("/health")
def health() -> dict[str, str]:
    return {"status": "ok", "version": "0.6.0"}


@app.post("/api/upload")
async def upload(files: list[UploadFile] = File(...)) -> dict[str, Any]:
    if not files:
        raise HTTPException(status_code=400, detail="Choose at least one file.")
    if len(files) > MAX_FILES:
        raise HTTPException(status_code=400, detail=f"Maximum {MAX_FILES} files per run.")

    summaries: list[dict[str, Any]] = []
    all_datasets: list[dict[str, Any]] = []
    parsed_count = 0

    for upload_file in files:
        filename = Path(upload_file.filename or "unnamed").name
        content = await upload_file.read()
        if len(content) > MAX_FILE_BYTES:
            summaries.append({"name": filename, "ok": False, "error": "File exceeds 25 MB."})
            continue
        try:
            parsed = parse_uploaded_file(filename, content)
            parsed_count += 1
            datasets = parsed.get("datasets", [])
            all_datasets.extend(datasets)
            file_signatures = {tuple(dataset.get("signature", [])) for dataset in datasets}
            file_tables_differ = len(file_signatures) > 1
            file_tables_merged = len(datasets) > 1 and len(file_signatures) == 1
            dataset_priority = {"promotion_table": 100, "student_roster": 95, "monitoring_table": 80, "assessment_results": 70, "financial_table": 40, "generic_table": 30, "unknown_table": 10}
            file_primary = max(
                datasets,
                key=lambda dataset: (dataset_priority.get(dataset.get("dataset_type"), 0), len(dataset.get("records", []))),
                default=None,
            )
            summaries.append({
                "name": filename,
                "ok": True,
                "detected_type": parsed.get("detected_type", "Document"),
                "source_kind": parsed.get("source_kind", "document"),
                "document_type": parsed.get("document_type", "unknown_document"),
                "document_confidence": parsed.get("document_confidence", 0),
                "classification_reasons": parsed.get("classification_reasons", []),
                "dataset_count": len(datasets),
                "dataset_types": sorted({dataset.get("dataset_type", "unknown_table") for dataset in datasets}),
                "rows": sum(len(dataset.get("records", [])) for dataset in datasets),
                "columns": max((len(dataset.get("columns", [])) for dataset in datasets), default=0),
                "table_count": parsed.get("table_count"),
                "usable_table_count": parsed.get("usable_table_count", len(datasets)),
                "sheet_count": parsed.get("sheet_count"),
                "warnings": parsed.get("warnings", []),
                "table_summaries": parsed.get("table_summaries", []),
                "tables_differ": file_tables_differ,
                "tables_merged": file_tables_merged,
                "selected_table": (file_primary.get("table_index") if file_primary else None),
                "warning_code": ("different_word_tables" if file_tables_differ else "similar_word_tables_merged" if file_tables_merged else None),
                "header_mode": (datasets[0].get("header_mode") if datasets else "none"),
                "header_rows": (datasets[0].get("header_rows", 1) if datasets else 0),
                "continuation_rows_merged": sum(dataset.get("diagnostics", {}).get("continuation_rows_merged", 0) for dataset in datasets),
                "inferred_class": ((datasets[0].get("class_info") or {}).get("value") if datasets else None),
                "class_source": ((datasets[0].get("class_info") or {}).get("source") if datasets else None),
                "class_conflict": ((datasets[0].get("class_info") or {}).get("conflict", False) if datasets else False),
                "excerpt": parsed.get("document_text_excerpt", ""),
            })
        except Exception as exc:
            summaries.append({"name": filename, "ok": False, "error": str(exc)})

    if not parsed_count:
        raise HTTPException(status_code=400, detail={"message": "No files could be parsed.", "files": summaries})

    groups = group_datasets(all_datasets)
    session_id = secrets.token_urlsafe(18)
    group_map = {group["id"]: group for group in groups}
    active = groups[0] if groups else None
    SESSIONS[session_id] = {
        "columns": active["columns"] if active else [],
        "records": active["records"] if active else [],
        "files": summaries,
        "dataset_groups": group_map,
        "active_dataset_id": active["id"] if active else None,
    }
    if len(SESSIONS) > 30:
        oldest = next(iter(SESSIONS))
        if oldest != session_id:
            SESSIONS.pop(oldest, None)

    public_groups = [
        {key: value for key, value in group.items() if key not in {"columns", "records"}}
        for group in groups
    ]
    return {
        "session_id": session_id,
        "files": summaries,
        "dataset_groups": public_groups,
        "active_dataset_id": active["id"] if active else None,
        "columns": active["columns"] if active else [],
        "row_count": active["row_count"] if active else 0,
        "message": "No reliable data tables were found." if not active else None,
    }


@app.post("/api/select-dataset")
def select_dataset(request: SelectDatasetRequest) -> dict[str, Any]:
    session = SESSIONS.get(request.session_id)
    if session is None:
        raise HTTPException(status_code=404, detail="Session expired. Upload the files again.")
    group = session.get("dataset_groups", {}).get(request.dataset_id)
    if group is None:
        raise HTTPException(status_code=404, detail="Dataset group was not found.")
    session["columns"] = group["columns"]
    session["records"] = group["records"]
    session["active_dataset_id"] = group["id"]
    return {
        "active_dataset_id": group["id"],
        "columns": group["columns"],
        "row_count": group["row_count"],
        "dataset_type": group["dataset_type"],
        "diagnostics": group.get("diagnostics", {}),
    }


@app.post("/api/preview")
def preview(request: PreviewRequest) -> dict[str, Any]:
    session = SESSIONS.get(request.session_id)
    if session is None:
        raise HTTPException(status_code=404, detail="Session expired. Upload the files again.")
    rows = build_output_rows(session, request)
    headers = list(rows[0].keys()) if rows else []
    return {
        "headers": headers,
        "rows": rows[: request.limit],
        "row_count": len(rows),
        "preview_count": min(len(rows), request.limit),
    }


@app.post("/api/export")
def export(request: ExportRequest) -> Response:
    session = SESSIONS.get(request.session_id)
    if session is None:
        raise HTTPException(status_code=404, detail="Session expired. Upload the files again.")
    rows = build_output_rows(session, request)
    if not rows:
        raise HTTPException(status_code=400, detail="No rows available for export.")

    export_format = request.format.lower()
    active_group = session.get("dataset_groups", {}).get(session.get("active_dataset_id"), {})
    base = safe_filename(f"FormaFlow_{active_group.get('dataset_type', 'Output')}")
    if export_format == "xlsx":
        content = create_xlsx(rows)
        media_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        filename = f"{base}.xlsx"
    elif export_format == "csv":
        content = create_csv(rows)
        media_type = "text/csv; charset=utf-8"
        filename = f"{base}.csv"
    elif export_format == "docx":
        content = create_docx(rows)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        filename = f"{base}.docx"
    else:
        raise HTTPException(status_code=400, detail="Format must be xlsx, csv, or docx.")

    return Response(
        content=content,
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


if __name__ == "__main__":
    import uvicorn

    uvicorn.run("app:app", host="127.0.0.1", port=8000, reload=False)
